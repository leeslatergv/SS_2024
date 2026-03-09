"""
Build the Word document version of the Staff Standards FA report.
Uses python-docx with embedded charts.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.expanduser("~/nhs-staff-standards/output")
REPORT_DIR = os.path.expanduser("~/nhs-staff-standards/report")
os.makedirs(REPORT_DIR, exist_ok=True)

# DHSC / GOV.UK colours
DHSC_TEAL = RGBColor(0x00, 0xA9, 0x90)
DHSC_DARK_TEAL = RGBColor(0x00, 0x86, 0x74)
DHSC_DEEPEST = RGBColor(0x00, 0x5A, 0x4E)
GOVUK_DARK_GREY = RGBColor(0x50, 0x5A, 0x5F)
GOVUK_BLACK = RGBColor(0x0B, 0x0C, 0x0C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GOVUK_LIGHT_GREY = RGBColor(0xF3, 0xF2, 0xF1)


def set_cell_shading(cell, color_hex):
    """Set background colour of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def add_formatted_table(doc, headers, rows, col_widths=None, header_color="00A990"):
    """Add a formatted table with NHS-style header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = GOVUK_BLACK
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F5")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def build_report():
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Styles ──────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    style.font.color.rgb = GOVUK_BLACK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, (size, color) in enumerate([(22, DHSC_DEEPEST), (16, DHSC_TEAL),
                                            (13, DHSC_DEEPEST), (11, GOVUK_DARK_GREY)], 1):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Arial"
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = True
        h.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
        h.paragraph_format.space_after = Pt(6)

    # ── Title page ──────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Measuring the NHS Staff Standards")
    run.font.size = Pt(28)
    run.font.color.rgb = DHSC_DEEPEST
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Factor Analysis of the NHS Staff Survey")
    run.font.size = Pt(16)
    run.font.color.rgb = DHSC_TEAL

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Analytical Report\nDHSC Social Research & Economics\nFebruary 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = GOVUK_DARK_GREY

    doc.add_page_break()

    # ── Executive Summary ───────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        "This report assesses whether the six NHS Staff Standards can be reliably "
        "measured using questions from the NHS Staff Survey, and recommends how to "
        "score and compare them across trusts."
    )

    doc.add_paragraph(
        "We conducted exploratory factor analysis (EFA) and confirmatory factor analysis "
        "(CFA) on trust-level data from 210 NHS trusts across four survey years (2021\u20132024), "
        "using 50 candidate survey questions spanning all six standards."
    )

    doc.add_heading("Key findings", level=3)

    findings = [
        "The six standards do not correspond to six separate empirical dimensions. "
        "The data supports three underlying constructs at trust level, with a possible fourth (appraisal quality).",
        "Three standards (Flexible Working, Effective Line Management, Supportive Environment) are "
        "empirically indistinguishable at trust level, with latent factor correlations of "
        "r = 0.91 to 0.99. Trusts that perform well on one perform well on all three.",
        "Violence and Sexual Safety also merge into a single \u2018patient-facing harm\u2019 "
        "dimension (r = 0.96).",
        "Tackling Racism is the most distinct standard, with moderate correlations to others (r = 0.58\u20130.87).",
        "Despite this, 36% of variation is standard-specific. Trusts do show meaningfully "
        "different profiles across standards, and the six-standard framework has monitoring value.",
    ]
    for f in findings:
        doc.add_paragraph(f, style="List Bullet")

    # Structure diagram
    doc.add_paragraph()
    doc.add_picture(os.path.join(OUTPUT_DIR, "report_structure_comparison.png"), width=Inches(6))
    cap = doc.add_paragraph("Figure 1: Theorised six-standard structure (left) vs empirical three-cluster structure (right)")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GOVUK_DARK_GREY

    doc.add_heading("Recommendations", level=3)
    recs = [
        "Use the recommended question sets per standard (Section 7).",
        "Score each standard as a simple average of its items after direction-correction.",
        "Compare trusts using z-score profiles rather than raw scores.",
        "Report the standards as related but complementary lenses on trust performance, "
        "not as six independent dimensions.",
    ]
    for r in recs:
        doc.add_paragraph(r, style="List Number")

    doc.add_page_break()

    # ── 1. Background ──────────────────────────────────────────────
    doc.add_heading("1. Background and Purpose", level=1)

    doc.add_paragraph(
        "DHSC has set out six standards for NHS employers:"
    )

    add_formatted_table(doc,
        ["#", "Standard", "Sub-standards"],
        [
            ["1", "Reducing Violence Against Staff", "\u2014"],
            ["2", "Tackling Racism", "\u2014"],
            ["3", "Improving Sexual Safety", "\u2014"],
            ["4", "Promoting Flexible Working", "\u2014"],
            ["5", "Effective Line Management", "Line management; Appraisal; Development"],
            ["6", "Creating a Supportive Environment", "Facilities; Support for Wellbeing"],
        ],
        col_widths=[1.5, 5.5, 7],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "To operationalise these standards for trust-level monitoring, two questions must be answered:"
    )

    p = doc.add_paragraph()
    run = p.add_run("A. Which NHS Staff Survey questions should measure each standard? ")
    run.bold = True
    p.add_run(
        "The survey contains over 100 questions. Some map directly to standards "
        "(e.g. q13a on physical violence maps to Standard 1). Others are less clear-cut."
    )

    p = doc.add_paragraph()
    run = p.add_run("B. How should scores be aggregated, and can standards be meaningfully compared? ")
    run.bold = True
    p.add_run(
        "This requires understanding whether the standards are empirically distinct "
        "or substantially overlapping, and what scoring methodology best captures trust-level variation."
    )

    doc.add_paragraph(
        "Factor analysis is the appropriate tool for both questions. It reveals the latent "
        "structure underlying observed survey responses and tests whether the theorised "
        "six-standard structure matches the data."
    )

    # ── 2. Data ────────────────────────────────────────────────────
    doc.add_heading("2. Data", level=1)

    doc.add_heading("2.1 Source", level=2)
    doc.add_paragraph(
        "NHS Staff Survey benchmark data (trust-level aggregates) for 2021, 2022, 2023 and 2024. "
        "Each observation is a trust-year, with variables representing the proportion of staff at that "
        "trust giving a particular response."
    )

    doc.add_heading("2.2 Sample", level=2)
    doc.add_paragraph(
        "210 NHS trusts across five trust types, yielding 821 trust-year observations. "
        "For cross-sectional analysis, we averaged each variable across years per trust, "
        "producing 210 trust-level means."
    )

    add_formatted_table(doc,
        ["Trust type", "Number of trusts"],
        [
            ["Acute & Acute Community Trusts", "122"],
            ["Mental Health, Learning Disability & Community Trusts", "50"],
            ["Acute Specialist Trusts", "13"],
            ["Community Trusts", "14"],
            ["Ambulance Trusts", "11"],
            ["Total", "210"],
        ],
        col_widths=[10, 4],
    )

    doc.add_heading("2.3 Variables", level=2)
    doc.add_paragraph(
        "We selected 50 variables with face-valid relevance to at least one standard. "
        "This is deliberately broader than any initial proposed question set to avoid the "
        "circularity of only testing pre-assigned questions. Variables are proportions "
        "bounded between 0 and 1, not individual-level Likert responses."
    )

    # ── 3. Methodology ─────────────────────────────────────────────
    doc.add_heading("3. Methodology", level=1)

    doc.add_heading("3.1 Analytical strategy", level=2)
    doc.add_paragraph("We followed a three-phase approach:")
    phases = [
        "Exploratory Factor Analysis (EFA) on the full 50-item pool to identify the "
        "empirical factor structure without imposing the six-standard framework.",
        "Confirmatory Factor Analysis (CFA) comparing competing models: one-factor, "
        "three-factor (empirically merged), and six-factor (theory-driven).",
        "Profile analysis to assess whether trusts show meaningful differentiation "
        "across standards, even where standards are highly correlated.",
    ]
    for i, ph in enumerate(phases, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Phase {i}: ")
        run.bold = True
        p.add_run(ph)

    doc.add_heading("3.2 EFA specification", level=2)
    specs = [
        ("Extraction method:", "Minimum residual (MinRes), robust to non-normal data."),
        ("Rotation:", "Promax (oblique), allowing factors to correlate."),
        ("Number of factors:", "Determined by Horn\u2019s parallel analysis (500 iterations, 95th percentile)."),
        ("Adequacy tests:", "Kaiser-Meyer-Olkin (KMO) measure and Bartlett\u2019s test of sphericity."),
    ]
    for label, desc in specs:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("3.3 CFA specification", level=2)
    doc.add_paragraph(
        "CFA models were estimated using maximum likelihood (ML) in R\u2019s lavaan package "
        "(version 0.6-21). To address extreme multicollinearity in the manager quality items "
        "(q9a\u2013i, which correlate at r = 0.95\u20130.996 at trust level), these nine items were "
        "composited into a single \u2018manager quality\u2019 variable for CFA. Similarly, three "
        "interpersonal culture items (q8b\u2013d, r = 0.99) were composited. This reduced CFA "
        "items from 50 to 20, giving an N:p ratio of 10.5:1."
    )

    doc.add_heading("3.4 Important caveats", level=2)

    caveats = [
        ("Ecological data.", "We are analysing trust-level aggregates, not individual responses. "
         "Factor structures at the aggregate level may differ from individual-level structures. "
         "Our findings describe the structure of trust-level differences, which is appropriate "
         "for trust-level monitoring but should not be interpreted as evidence about "
         "individual-level measurement properties."),
        ("Small sample.", "N = 210 trusts is adequate for EFA (KMO = 0.96) but marginal for "
         "complex CFA models. We focus on relative model comparisons rather than absolute "
         "fit benchmarks."),
        ("Bounded proportions.", "Variables are constrained to [0, 1] with floor/ceiling "
         "effects at some trusts. This can attenuate correlations and distort factor solutions."),
    ]
    for label, desc in caveats:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.bold = True
        run.italic = True
        p.add_run(desc)

    doc.add_page_break()

    # ── 4. EFA Results ─────────────────────────────────────────────
    doc.add_heading("4. Results: Exploratory Factor Analysis", level=1)

    doc.add_heading("4.1 Adequacy and number of factors", level=2)

    p = doc.add_paragraph()
    p.add_run("KMO = 0.960 ").bold = True
    p.add_run("(superb; above 0.90 is considered excellent). ")
    p.add_run("Bartlett\u2019s test: ").bold = True
    p.add_run("\u03c7\u00b2 = 27,123, p < 0.001. The data is highly suitable for factor analysis.")

    doc.add_paragraph(
        "Parallel analysis suggested four factors. The first factor is dominant, accounting "
        "for 66% of the eigenvalue sum, indicating a strong general factor underlying all items."
    )

    # Scree plot
    doc.add_picture(os.path.join(OUTPUT_DIR, "scree_expanded.png"), width=Inches(5.5))
    cap = doc.add_paragraph("Figure 2: Scree plot with parallel analysis threshold (50 items, N = 210)")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GOVUK_DARK_GREY

    doc.add_heading("4.2 Four-factor solution", level=2)
    doc.add_paragraph(
        "The promax-rotated four-factor solution explained 76% of total variance:"
    )

    add_formatted_table(doc,
        ["Factor", "Interpretation", "Key items loading"],
        [
            ["F1", "Good immediate management\nand work environment",
             "All manager items (q9a\u2013i), flexible working (q4d, q6d), "
             "interpersonal respect (q7c, q8b\u2013d), strained relationships (q5c)"],
            ["F2", "Organisational culture\nand fairness",
             "Fair career progression (q15), respects differences (q21), "
             "safe to speak up (q25e), wellbeing action (q11a), "
             "materials/equipment (q3h), career development (q24b\u2013e)"],
            ["F3", "Patient-facing harm",
             "Violence from patients (q13a), harassment from patients (q14a), "
             "sexual behaviour from patients (q17a) and staff (q17b)"],
            ["F4", "Appraisal quality",
             "Appraisal improved job (q23b), clear objectives (q23c), "
             "felt valued (q23d)"],
        ],
        col_widths=[1.5, 4.5, 10],
    )

    doc.add_paragraph()

    # Heatmap
    doc.add_picture(os.path.join(OUTPUT_DIR, "heatmap_4factor_expanded.png"), width=Inches(5.5))
    cap = doc.add_paragraph("Figure 3: EFA pattern matrix heatmap (4-factor promax solution)")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GOVUK_DARK_GREY

    doc.add_heading("4.3 Mapping EFA factors to standards", level=2)

    add_formatted_table(doc,
        ["EFA Factor", "Standards captured"],
        [
            ["F1: Good management", "Flexible Working + Line Management (line management) + Supportive Environment (culture)"],
            ["F2: Org culture", "Tackling Racism + Supportive Environment (wellbeing, facilities) + Line Management (development)"],
            ["F3: Patient harm", "Reducing Violence + Sexual Safety"],
            ["F4: Appraisal", "Line Management (appraisal only)"],
        ],
        col_widths=[4, 12],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Key observation: ")
    run.bold = True
    p.add_run(
        "When forced to six factors, the solution did not produce six clean factors "
        "matching the theorised standards. The additional factors were weakly defined "
        "with cross-loadings and Heywood cases (communalities exceeding 1.0)."
    )

    doc.add_page_break()

    # ── 5. CFA Results ─────────────────────────────────────────────
    doc.add_heading("5. Results: Confirmatory Factor Analysis", level=1)

    doc.add_heading("5.1 Model comparison", level=2)

    # CFA comparison chart
    doc.add_picture(os.path.join(OUTPUT_DIR, "report_cfa_comparison.png"), width=Inches(4.5))
    cap = doc.add_paragraph("Figure 4: Comparative Fit Index (CFI) across competing CFA models")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GOVUK_DARK_GREY

    doc.add_paragraph()

    add_formatted_table(doc,
        ["Model", "\u03c7\u00b2", "df", "\u03c7\u00b2/df", "CFI", "TLI", "RMSEA", "SRMR"],
        [
            ["A. Single factor", "3,269", "170", "19.2", "0.517", "0.460", "0.295", "0.135"],
            ["B. Three-factor (merged)", "2,928", "167", "17.5", "0.569", "0.510", "0.281", "0.182"],
            ["D. Six-factor (theory)", "2,811", "155", "18.1", "0.586", "0.492", "0.286", "0.182"],
        ],
        col_widths=[5, 1.5, 1, 1.5, 1.5, 1.5, 1.5, 1.5],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "No model achieved conventional thresholds for acceptable fit (CFI > 0.90, "
        "RMSEA < 0.08). This is expected given ecological data, small N, and bounded "
        "proportions (see Section 3.4). However, the relative comparison is informative: "
        "the six-factor model has the lowest \u03c7\u00b2 and highest CFI, suggesting the "
        "finer-grained structure adds value beyond a single general factor."
    )

    doc.add_heading("5.2 Factor correlations (discriminant validity)", level=2)
    doc.add_paragraph(
        "The critical finding comes from the latent factor correlations estimated by "
        "the six-factor CFA model:"
    )

    doc.add_picture(os.path.join(OUTPUT_DIR, "report_factor_correlations.png"), width=Inches(5))
    cap = doc.add_paragraph(
        "Figure 5: Latent factor correlations from 6-factor CFA. "
        "Bold values indicate r > 0.90 (empirically redundant). "
        "Red box highlights the FlexWorking/LineMgmt/SupportEnv cluster. "
        "Dashed circles highlight the Violence/SexSafety pair."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GOVUK_DARK_GREY

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Interpretation of key pairs:")
    run.bold = True

    pairs = [
        ("LineMgmt \u2194 SupportEnv (r = 0.989):", "Statistically indistinguishable. "
         "A trust\u2019s line management score predicts its supportive environment score "
         "with near-perfect accuracy."),
        ("FlexWorking \u2194 LineMgmt (r = 0.976):", "Also effectively identical. "
         "Whether staff report good flexible working is almost entirely explained by "
         "whether they report good line management."),
        ("Violence \u2194 SexSafety (r = 0.955):", "Trusts with violence problems "
         "have sexual safety problems and vice versa."),
        ("Racism \u2194 SupportEnv (r = 0.865):", "High but not redundant. "
         "Discrimination is correlated with poor organisational culture but retains "
         "a meaningful unique component."),
        ("Racism \u2194 SexSafety (r = \u22120.582):", "The most distinct pair. "
         "These measure genuinely different things."),
    ]
    for label, desc in pairs:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(label + " ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "In measurement terms, correlations above 0.90 between latent factors indicate "
        "that the constructs cannot be empirically distinguished. The covariance matrix "
        "of latent variables was flagged as non-positive-definite by lavaan, confirming "
        "that the six-factor structure is empirically untenable as six separate dimensions."
    )

    doc.add_heading("5.3 Standardised factor loadings", level=2)
    doc.add_paragraph(
        "Despite the redundancy between factors, items loaded strongly on their assigned factors:"
    )

    doc.add_picture(os.path.join(OUTPUT_DIR, "report_loadings.png"), width=Inches(6))
    cap = doc.add_paragraph("Figure 6: CFA standardised loadings by standard")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GOVUK_DARK_GREY

    doc.add_paragraph(
        "All loadings are statistically significant (p < 0.001). The weakest is "
        "appraisal_improved_job (0.276), suggesting appraisal quality captures something "
        "partially distinct from broader line management."
    )

    doc.add_page_break()

    # ── 6. Profile Analysis ────────────────────────────────────────
    doc.add_heading("6. Results: Profile Analysis", level=1)

    doc.add_heading("6.1 Variance decomposition", level=2)

    doc.add_picture(os.path.join(OUTPUT_DIR, "report_variance_decomposition.png"), width=Inches(3.5))
    cap = doc.add_paragraph("Figure 7: Decomposition of trust-level variance across standards")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GOVUK_DARK_GREY

    doc.add_paragraph()
    doc.add_paragraph(
        "The mean inter-standard correlation is r = 0.80, implying that 64% of "
        "trust-level variation is a general \u2018good trust\u2019 factor and 36% is "
        "standard-specific. Most variation is shared, but a meaningful minority is not."
    )

    doc.add_heading("6.2 Trust profiles", level=2)
    doc.add_paragraph(
        "To assess whether trusts show meaningfully different patterns (not just levels), "
        "we computed within-trust z-score profiles across the five standards. "
        "The mean profile variability was 0.40 SD, indicating meaningful differentiation."
    )

    doc.add_picture(os.path.join(OUTPUT_DIR, "report_trust_profiles.png"), width=Inches(6.2))
    cap = doc.add_paragraph(
        "Figure 8: Example trust z-score profiles. Trust RNK, for instance, "
        "scores above average on violence (low violence) but nearly three standard "
        "deviations below average on racism."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GOVUK_DARK_GREY

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Key implication: ")
    run.bold = True
    p.add_run(
        "The six-standard framework has monitoring value because it directs attention "
        "to specific areas where a trust may underperform relative to its general level. "
        "A single composite score would mask these patterns."
    )

    # ── 6.3 Sub-standard structure ─────────────────────────────────
    doc.add_heading("6.3 Sub-standard structure", level=2)

    doc.add_paragraph(
        "Two standards have defined sub-standards: Effective Line Management "
        "(line management, appraisal, development) and Supportive Environment "
        "(facilities, wellbeing). The factor analysis has something to say about "
        "whether these sub-groupings hold in the data."
    )

    doc.add_heading("Effective Line Management", level=3)

    elm_findings = [
        ("Line management (q9a\u2013i).",
         "All nine manager quality items correlate at r = 0.95 to 0.996 at trust level. "
         "They are a single, undifferentiated block. There is no empirical basis for "
         "distinguishing between, say, \u2018manager gives feedback\u2019 and "
         "\u2018manager shows interest in wellbeing\u2019 at trust level."),
        ("Appraisal (q23a\u2013d).",
         "Emerged as a distinct sub-factor (Factor 4) in the EFA. This is one of the "
         "clearest findings: having a good line manager does not automatically translate "
         "into a useful appraisal. The sub-standard distinction is empirically validated."),
        ("Development (q24b, q24d, q24e).",
         "In the EFA, development items loaded more strongly on the organisational "
         "culture factor (F2) than on the manager/team factor (F1). This suggests that "
         "access to learning and career development is driven more by organisational "
         "culture and investment than by individual line manager quality. The data would "
         "place development closer to Supportive Environment than to Line Management."),
    ]
    for label, desc in elm_findings:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(label + " ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("Supportive Environment", level=3)

    se_findings = [
        ("Facilities (q3h, q3i).",
         "Loaded on the organisational culture factor (F2). Whether a trust has "
         "adequate materials and enough staff is an institutional-level decision, "
         "consistent with treating this as a distinct sub-standard."),
        ("Wellbeing (q8b\u2013d, q11a).",
         "This sub-standard is split across two empirical factors. The interpersonal "
         "culture items (q8b\u2013d: kindness, respect, appreciation) loaded with the "
         "manager/team factor (F1), suggesting they reflect team-level culture driven "
         "by line management. The organisational wellbeing item (q11a) loaded with the "
         "organisational culture factor (F2). In other words, \u2018people being kind to "
         "each other\u2019 is a team phenomenon, while \u2018the organisation takes action "
         "on wellbeing\u2019 is an institutional one."),
    ]
    for label, desc in se_findings:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(label + " ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Summary: ")
    run.bold = True
    p.add_run(
        "The appraisal sub-standard within ELM is clearly validated by the data. "
        "The development sub-standard may sit more naturally under Supportive Environment. "
        "The wellbeing sub-standard straddles two empirical factors, with interpersonal "
        "culture being team-driven and organisational wellbeing being institution-driven."
    )

    doc.add_page_break()

    # ── 7. Recommended Question Sets ───────────────────────────────
    doc.add_heading("7. Recommended Question Sets", level=1)

    doc.add_paragraph(
        "Based on the EFA loadings, CFA standardised loadings, and face validity, "
        "we recommend the following questions for each standard. Items are listed in "
        "order of loading strength."
    )

    # Standard 1
    doc.add_heading("Standard 1: Reducing Violence Against Staff", level=2)

    add_formatted_table(doc,
        ["Code", "Question", "Loading", "Direction"],
        [
            ["q14a", "Harassment/bullying/abuse from patients/public", "0.983", "Lower = better"],
            ["q13a", "Physical violence from patients/public", "0.903", "Lower = better"],
            ["q14b", "Harassment/bullying/abuse from managers", "0.614", "Lower = better"],
            ["q14c", "Harassment/bullying/abuse from colleagues", "0.468", "Lower = better"],
        ],
        col_widths=[1.5, 7, 1.8, 2.5],
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    p.add_run(
        "We include harassment items (q14a\u2013c) alongside physical violence. "
        "The reporting items (q13d, q14d) should be excluded from the composite as they "
        "measure reporting culture rather than incidence. They can be reported separately. "
        "Violence from managers (q13b) and colleagues (q13c) had near-zero variance at "
        "trust level and were dropped."
    )

    # Standard 2
    doc.add_heading("Standard 2: Tackling Racism", level=2)

    add_formatted_table(doc,
        ["Code", "Question", "Loading", "Direction"],
        [
            ["q15", "Fair career progression regardless of characteristics", "0.953", "Higher = better"],
            ["q21", "Organisation respects individual differences", "0.883", "Higher = better"],
            ["q16b", "Discrimination from managers/colleagues", "0.866", "Lower = better"],
            ["q16a", "Discrimination from patients/public", "0.723", "Lower = better"],
        ],
        col_widths=[1.5, 7, 1.8, 2.5],
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    p.add_run(
        "This is the most empirically distinct standard. q21 is not exclusively about "
        "race but loads strongly here because trusts with discrimination problems also "
        "score poorly on respecting differences."
    )

    # Standard 3
    doc.add_heading("Standard 3: Improving Sexual Safety", level=2)

    add_formatted_table(doc,
        ["Code", "Question", "Loading", "Direction"],
        [
            ["q17a", "Unwanted sexual behaviour from patients/public", "0.867", "Lower = better"],
            ["q17b", "Unwanted sexual behaviour from staff/colleagues", "0.849", "Lower = better"],
        ],
        col_widths=[1.5, 7, 1.8, 2.5],
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    p.add_run(
        "Only two items available (introduced 2023). Empirically very close to "
        "Reducing Violence (CFA r = 0.96). Could be monitored jointly as part of "
        "a broader \u2018Staff Safety\u2019 construct."
    )

    # Standard 4
    doc.add_heading("Standard 4: Promoting Flexible Working", level=2)

    add_formatted_table(doc,
        ["Code", "Question", "Loading", "Direction"],
        [
            ["q6c", "I achieve a good balance between my work life and home life", "0.972", "Higher = better"],
            ["q4d", "Satisfaction with opportunities for flexible working", "0.927", "Higher = better"],
        ],
        col_widths=[1.5, 7, 1.8, 2.5],
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    p.add_run(
        "Empirically very close to Line Management (CFA r = 0.98). "
        "Trusts with good managers automatically score well on flexible working. "
        "Can be retained as a separate monitoring lens but should not be interpreted as "
        "measuring a distinct dimension."
    )

    # Standard 5
    doc.add_heading("Standard 5: Effective Line Management", level=2)

    doc.add_heading("5a. Line management", level=3)
    add_formatted_table(doc,
        ["Code", "Question", "Direction"],
        [
            ["q9d", "Manager takes positive interest in health/wellbeing", "Higher = better"],
            ["q9f", "Manager works together to understand problems", "Higher = better"],
            ["q9i", "Manager takes effective action to help", "Higher = better"],
        ],
        col_widths=[1.5, 8, 2.5],
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    p.add_run(
        "All nine q9 items correlate at r = 0.95 to 0.996 at trust level. "
        "They measure the same thing. Either use all nine and average, or select "
        "three representative items as above. Results are identical to three decimal places."
    )

    doc.add_heading("5b. Appraisal", level=3)
    add_formatted_table(doc,
        ["Code", "Question", "Loading", "Direction"],
        [
            ["q23a", "Had appraisal in last 12 months", "0.574", "Higher = better"],
            ["q23b", "Appraisal helped improve how I do my job", "0.276", "Higher = better"],
            ["q23c", "Appraisal helped agree clear objectives", "\u2014", "Higher = better"],
            ["q23d", "Appraisal left me feeling valued", "\u2014", "Higher = better"],
        ],
        col_widths=[1.5, 7, 1.8, 2.5],
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    p.add_run(
        "Appraisal quality emerged as a distinct sub-factor in EFA. "
        "Having a good manager does not automatically translate into a useful appraisal. "
        "Track separately within this standard."
    )

    doc.add_heading("5c. Development", level=3)
    add_formatted_table(doc,
        ["Code", "Question", "Loading", "Direction"],
        [
            ["q24e", "Access to right L&D opportunities", "0.875", "Higher = better"],
            ["q24b", "Opportunities for career development", "\u2014", "Higher = better"],
            ["q24d", "Feel supported to develop potential", "\u2014", "Higher = better"],
        ],
        col_widths=[1.5, 7, 1.8, 2.5],
    )

    # Standard 6
    doc.add_heading("Standard 6: Creating a Supportive Environment", level=2)

    doc.add_heading("6a. Facilities", level=3)
    add_formatted_table(doc,
        ["Code", "Question", "Loading", "Direction"],
        [
            ["q3h", "Adequate materials, supplies and equipment", "0.804", "Higher = better"],
            ["q3i", "Enough staff to do my job properly", "0.699", "Higher = better"],
            ["q22", "Can eat nutritious affordable food at work", "\u2014", "Higher = better"],
        ],
        col_widths=[1.5, 7, 1.8, 2.5],
    )

    doc.add_heading("6b. Support for wellbeing", level=3)
    add_formatted_table(doc,
        ["Code", "Question", "Loading", "Direction"],
        [
            ["q11a", "Organisation takes positive action on wellbeing", "0.925", "Higher = better"],
            ["q8b", "People are understanding and kind to one another", "\u2014", "Higher = better"],
            ["q8c", "People are polite and treat each other with respect", "\u2014", "Higher = better"],
            ["q8d", "People show appreciation to one another", "\u2014", "Higher = better"],
        ],
        col_widths=[1.5, 7, 1.8, 2.5],
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    p.add_run(
        "q8b\u2013d are highly collinear (r = 0.99) and can be composited into a "
        "single \u2018interpersonal culture\u2019 score."
    )

    doc.add_page_break()

    # ── 7.5 Cross-validation ──────────────────────────────────────
    doc.add_heading("8. Cross-Validation", level=1)

    doc.add_paragraph(
        "To test the robustness of these findings, we replicated the analysis using "
        "a temporal hold-out design: EFA on 2023 data to discover the structure, "
        "CFA on independent 2024 data to test it. Both years are post-COVID. "
        "This mirrors the methodology used for the People Promise analysis."
    )

    doc.add_heading("8.1 EFA on 2023 (exploration sample)", level=2)
    cv_efa = [
        "N = 207 trusts. KMO = 0.961.",
        "Parallel analysis: 4 factors (same as main analysis).",
        "Same factor structure: F1 = manager/team, F2 = org culture, "
        "F3 = patient-facing harm, F4 = appraisal quality.",
        "Variance explained: 73.4%.",
    ]
    for item in cv_efa:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("8.2 CFA on 2024 (confirmation sample)", level=2)
    doc.add_paragraph(
        "We tested the same three models on the independent 2024 data (N = 210):"
    )

    add_formatted_table(doc,
        ["Model", "\u03c7\u00b2", "df", "\u03c7\u00b2/df", "CFI", "TLI", "RMSEA", "SRMR"],
        [
            ["A. 6-factor (theory)", "2,493", "155", "16.1", "0.603", "0.513", "0.268", "0.178"],
            ["B. 3-factor (merged)", "2,647", "167", "15.9", "0.579", "0.521", "0.266", "0.176"],
            ["C. 1-factor", "2,997", "170", "17.6", "0.520", "0.464", "0.281", "0.136"],
        ],
        col_widths=[5, 1.5, 1, 1.5, 1.5, 1.5, 1.5, 1.5],
    )

    doc.add_paragraph()
    doc.add_heading("8.3 Factor correlations on hold-out data", level=2)

    doc.add_paragraph(
        "The latent factor correlations from the 6-factor CFA on independent 2024 data "
        "closely replicate the main analysis:"
    )

    add_formatted_table(doc,
        ["Factor pair", "Main analysis\n(all years)", "2024 hold-out", "Interpretation"],
        [
            ["FlexWorking \u2194 LineMgmt", "0.976", "0.965", "Redundant"],
            ["LineMgmt \u2194 SupportEnv", "0.989", "1.017*", "Redundant"],
            ["Violence \u2194 SexSafety", "0.955", "0.932", "Redundant"],
            ["FlexWorking \u2194 SupportEnv", "0.914", "0.886", "High"],
            ["Racism \u2194 SexSafety", "\u22120.582", "\u22120.520", "Distinct"],
        ],
        col_widths=[4.5, 3, 3, 3],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("*")
    run.bold = True
    p.add_run(
        "A correlation exceeding 1.0 is mathematically impossible and indicates that "
        "the model cannot distinguish these two factors at all. lavaan flagged the "
        "covariance matrix as non-positive-definite, confirming the redundancy."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Conclusion: ")
    run.bold = True
    p.add_run(
        "All findings replicate on independent data. The three-cluster structure, "
        "the redundancy of FlexWorking/LineMgmt/SupportEnv, and the distinctiveness "
        "of Tackling Racism are not artefacts of the averaging approach or of COVID-era data."
    )

    doc.add_page_break()

    # ── 9. Scoring Methodology ─────────────────────────────────────
    doc.add_heading("9. Scoring Methodology", level=1)

    doc.add_heading("9.1 Computing standard scores", level=2)
    doc.add_paragraph(
        "For each trust, each standard score is the simple average of its component "
        "items after direction-correction."
    )

    p = doc.add_paragraph()
    run = p.add_run("Direction correction: ")
    run.bold = True
    p.add_run(
        "Items where \u2018lower = better\u2019 (violence, harassment, discrimination, "
        "sexual behaviour) should be reversed by subtracting from 1. After reversal, "
        "all items are on a \u2018higher = better\u2019 scale and can be averaged."
    )

    doc.add_heading("9.2 Why simple averages", level=2)
    doc.add_paragraph(
        "We considered factor score estimation and loading-weighted averages but recommend "
        "simple averages because:"
    )
    reasons = [
        "They are transparent and reproducible without specialised software.",
        "With high within-standard correlations, the correlation between simple averages "
        "and factor scores is typically > 0.98.",
        "Weighting by loadings adds complexity without meaningful improvement in validity.",
    ]
    for r in reasons:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("9.3 Comparing trusts: z-score profiles", level=2)
    doc.add_paragraph(
        "Raw composite scores are not directly comparable across standards because "
        "standards have different means and variances."
    )

    p = doc.add_paragraph()
    run = p.add_run("Recommended approach: ")
    run.bold = True

    steps = [
        "Compute the composite score for each standard for each trust.",
        "Convert to z-scores: z = (score \u2212 mean) / SD across all trusts.",
        "Present as a profile showing relative strengths and weaknesses.",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")

    doc.add_paragraph(
        "A z-score of 0 means average performance. This allows meaningful statements "
        "like \u2018Trust X performs 1.5 SDs above average on Racism but 0.8 SDs below "
        "average on Violence.\u2019"
    )

    doc.add_heading("9.4 Interpreting cross-standard differences", level=2)

    interp = [
        "Rankings of trusts will be broadly similar across standards. "
        "A trust ranking 10th on Line Management will typically rank in the top 20 on "
        "most other standards.",
        "But not identical. The 36% standard-specific variance means some trusts will "
        "show distinctive profiles.",
        "Do not over-interpret small differences between highly correlated standards "
        "(e.g. Flexible Working vs Line Management). Differences of less than ~0.5 SD "
        "are likely noise.",
        "Do interpret differences between weakly correlated standards "
        "(e.g. Racism vs Sexual Safety). A trust scoring well on one and poorly on the "
        "other is showing a genuine pattern.",
    ]
    for item in interp:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ── 9. Summary ─────────────────────────────────────────────────
    doc.add_heading("10. Summary of Recommendations", level=1)

    doc.add_heading("On measurement (which questions)", level=2)
    m_recs = [
        "Use the question sets in Section 7.",
        "Exclude reporting items (q13d, q14d) from composites. Track them separately.",
        "Composite the nine manager items (q9a\u2013i) into a single score.",
        "Track appraisal quality separately within Standard 5.",
    ]
    for r in m_recs:
        doc.add_paragraph(r, style="List Number")

    doc.add_heading("On structure (how many standards)", level=2)
    s_recs = [
        "Keep six standards for policy communication.",
        "Acknowledge that three empirical clusters exist: Staff Safety (Violence + Sexual Safety), "
        "Management & Culture (Flexible Working + Line Management + Supportive Environment), "
        "and Organisational Fairness (Tackling Racism, most distinct).",
        "Consider whether Flexible Working and Supportive Environment need to be separate "
        "standards given they are empirically indistinguishable from Line Management.",
    ]
    for r in s_recs:
        doc.add_paragraph(r, style="List Number")

    doc.add_heading("On scoring and comparison", level=2)
    c_recs = [
        "Use simple averages of direction-corrected items.",
        "Compare trusts using z-score profiles, not raw scores.",
        "Report the 64/36 split (general vs specific variance) to set expectations.",
    ]
    for r in c_recs:
        doc.add_paragraph(r, style="List Number")

    # ── Save ────────────────────────────────────────────────────────
    output_path = os.path.join(REPORT_DIR, "Measuring_NHS_Staff_Standards_FA_Report.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


if __name__ == "__main__":
    build_report()
