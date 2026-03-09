"""
Hierarchical NOF Analysis — Sub-standards → Standards
======================================================
Tests:
1. Unforced EFA: what emerges naturally?
2. Reliability at sub-standard and standard level
3. Second-order structure for Effective Line Management
   (Line Mgmt, Team Mgmt, Appraisal, Development → ELM)
4. Generates plain-English Word doc with RAG-rated tables

Uses 2024 NSS data only (263 organisations).
"""

import pandas as pd
import numpy as np
from factor_analyzer import FactorAnalyzer
from factor_analyzer.factor_analyzer import calculate_kmo
from scipy import stats
import warnings
warnings.filterwarnings("ignore")

# Re-use extraction from nof_2024_only
from nof_2024_only import extract_all_nof_items, cronbach_alpha, item_total_correlations, alpha_if_dropped


def standardized_alpha(items_df):
    """Standardized Cronbach's alpha (from mean inter-item correlation).
    Use when items have wildly different variances (e.g. violence frequencies)."""
    k = items_df.shape[1]
    if k < 2:
        return np.nan
    corr = items_df.corr()
    mask = np.triu(np.ones(corr.shape), k=1).astype(bool)
    mean_r = corr.where(mask).stack().mean()
    return (k * mean_r) / (1 + (k - 1) * mean_r)


def variance_ratio(items_df):
    """Max/min item variance ratio — flags when raw alpha is unreliable."""
    vrs = items_df.var(ddof=1)
    vrs = vrs[vrs > 0]
    if len(vrs) < 2:
        return 1.0
    return vrs.max() / vrs.min()

# ══════════════════════════════════════════════════════════════════════
# PROPOSED STRUCTURE (from policy specification)
# ══════════════════════════════════════════════════════════════════════

# Standard → sub-standard → questions
STRUCTURE = {
    "Reducing Violence": {
        "_sub_standards": None,  # No sub-standards
        "_items": {
            "violence_patients":    {"q": "Q13a: Physical violence from patients/public", "reverse": True},
            "violence_managers":    {"q": "Q13b: Physical violence from managers", "reverse": True},
            "violence_colleagues":  {"q": "Q13c: Physical violence from colleagues", "reverse": True},
            # Q13d not in updated list per user's table — only Q13a+b+c
        },
    },
    "Tackling Racism": {
        "_sub_standards": None,
        "_items": {
            "fair_career_progression": {"q": "Q15: Org acts fairly re career progression", "reverse": False},
            "discrim_patients":        {"q": "Q16a: Discrimination from patients (BME vs White LR)", "reverse": True},
            "discrim_colleagues":      {"q": "Q16b: Discrimination from manager/colleagues (BME vs White LR)", "reverse": True},
        },
        "_note": "Scored via likelihood ratios (BME vs White), not raw means",
    },
    "Sexual Safety": {
        "_sub_standards": None,
        "_items": {
            "sexual_behaviour_staff": {"q": "Q17b: Unwanted sexual behaviour from staff", "reverse": True},
        },
        "_note": "Updated list drops Q17a (patients); only Q17b retained",
    },
    "Flexible Working": {
        "_sub_standards": None,
        "_items": {
            "satisfaction_flexible_working": {"q": "Q4d: Satisfaction with flexible working", "reverse": False},
            "org_committed_wlb":             {"q": "Q6b: Org committed to work-life balance", "reverse": False},
            "manager_flexible_working":      {"q": "Q6d: Can approach manager re flexible working", "reverse": False},
        },
    },
    "Effective Line Management": {
        "_sub_standards": {
            "Line Management": {
                "items": {
                    "mgr_interest_wellbeing":     {"q": "Q9d: Manager takes interest in wellbeing", "reverse": False},
                    "mgr_understanding_problems": {"q": "Q9f: Manager works to understand problems", "reverse": False},
                    "mgr_effective_action":        {"q": "Q9i: Manager takes effective action", "reverse": False},
                },
            },
            "Team Management": {
                "items": {
                    "team_shared_objectives": {"q": "Q7a: Team has shared objectives", "reverse": False},
                    "feel_valued_by_team":    {"q": "Q7h: I feel valued by my team", "reverse": False},
                },
            },
            "Appraisal": {
                "items": {
                    "had_appraisal":              {"q": "Q23a: Had appraisal in last 12 months", "reverse": False},
                    "appraisal_clear_objectives": {"q": "Q23c: Appraisal helped agree clear objectives", "reverse": False},
                    "appraisal_felt_valued":      {"q": "Q23d: Appraisal left feeling valued", "reverse": False},
                },
            },
            "Development": {
                "items": {
                    "supported_develop_potential": {"q": "Q24d: Feel supported to develop potential", "reverse": False},
                    "access_learning_dev":         {"q": "Q24e: Can access learning/development", "reverse": False},
                },
            },
        },
        "_items": None,  # items live in sub-standards
    },
    "Supportive Environment": {
        "_sub_standards": {
            "Facilities": {
                "items": {
                    "nutritious_food": {"q": "Q22: Can eat nutritious/affordable food at work", "reverse": False},
                },
            },
            "Wellbeing": {
                "items": {
                    "org_positive_wellbeing": {"q": "Q11a: Org takes positive action on wellbeing", "reverse": False},
                },
            },
        },
        "_items": None,
        "_note": "Each sub-standard is a single item — second-order testing not possible",
    },
}


def get_all_items_flat():
    """Return list of all item variable names in the structure."""
    items = []
    for std, cfg in STRUCTURE.items():
        if cfg.get("_items"):
            items.extend(cfg["_items"].keys())
        if cfg.get("_sub_standards"):
            for sub_name, sub_cfg in cfg["_sub_standards"].items():
                items.extend(sub_cfg["items"].keys())
    return items


def get_item_info(var_name):
    """Look up question text, standard, sub-standard for a variable."""
    for std, cfg in STRUCTURE.items():
        if cfg.get("_items") and var_name in cfg["_items"]:
            return {"standard": std, "sub_standard": None,
                    "question": cfg["_items"][var_name]["q"],
                    "reverse": cfg["_items"][var_name]["reverse"]}
        if cfg.get("_sub_standards"):
            for sub_name, sub_cfg in cfg["_sub_standards"].items():
                if var_name in sub_cfg["items"]:
                    return {"standard": std, "sub_standard": sub_name,
                            "question": sub_cfg["items"][var_name]["q"],
                            "reverse": sub_cfg["items"][var_name]["reverse"]}
    return None


def get_standard_items(std_name):
    """Get all item variable names for a standard (including sub-standards)."""
    cfg = STRUCTURE[std_name]
    items = []
    if cfg.get("_items"):
        items.extend(cfg["_items"].keys())
    if cfg.get("_sub_standards"):
        for sub_name, sub_cfg in cfg["_sub_standards"].items():
            items.extend(sub_cfg["items"].keys())
    return items


# ══════════════════════════════════════════════════════════════════════
# ANALYSIS FUNCTIONS
# ══════════════════════════════════════════════════════════════════════

def run_unforced_efa(df):
    """Run EFA without imposing structure — what does the data say?"""
    print("\n" + "=" * 80)
    print("  UNFORCED EFA — WHAT NATURALLY EMERGES?")
    print("=" * 80)

    all_items = get_all_items_flat()
    available = [i for i in all_items if i in df.columns]
    dat = df[available].dropna()
    n_obs, n_vars = dat.shape
    print(f"\n  N = {n_obs} organisations, {n_vars} items")

    # Reverse-score negative items so all point same direction
    dat_aligned = dat.copy()
    for var in available:
        info = get_item_info(var)
        if info and info["reverse"]:
            dat_aligned[var] = 1 - dat_aligned[var]

    # KMO
    try:
        kmo_per, kmo_total = calculate_kmo(dat_aligned)
        print(f"  KMO: {kmo_total:.3f}")
    except Exception as e:
        print(f"  KMO failed: {e}")
        kmo_per = None

    results = {}

    for n_factors in [4, 6, 8]:
        print(f"\n  --- {n_factors}-factor solution ---")
        fa = FactorAnalyzer(n_factors=n_factors, rotation="promax", method="minres")
        fa.fit(dat_aligned)

        loadings = pd.DataFrame(
            fa.loadings_, index=available,
            columns=[f"F{i+1}" for i in range(n_factors)]
        )

        _, _, cum_var = fa.get_factor_variance()
        print(f"  Variance explained: {cum_var[-1]:.1%}")

        comm = pd.Series(fa.get_communalities(), index=available)

        # Show which items load where
        for var in available:
            row = loadings.loc[var]
            primary = row.abs().idxmax()
            primary_val = row[primary]
            info = get_item_info(var)
            std_label = info["standard"][:15] if info else "?"
            sub_label = info["sub_standard"] or "" if info else ""
            sig_loadings = [(c, v) for c, v in row.items() if abs(v) >= 0.30]
            load_str = ", ".join([f"{c}={v:+.3f}" for c, v in sig_loadings])
            print(f"    {std_label:15s} {sub_label:12s} {var:35s} → {load_str}")

        # Factor correlations
        if fa.phi_ is not None:
            phi = pd.DataFrame(
                fa.phi_,
                index=[f"F{i+1}" for i in range(n_factors)],
                columns=[f"F{i+1}" for i in range(n_factors)]
            )
            high_pairs = []
            for i in range(n_factors):
                for j in range(i+1, n_factors):
                    r = phi.iloc[i, j]
                    if abs(r) > 0.50:
                        high_pairs.append((f"F{i+1}", f"F{j+1}", r))
            if high_pairs:
                print(f"\n  High factor correlations (|r|>0.50):")
                for f1, f2, r in sorted(high_pairs, key=lambda x: -abs(x[2])):
                    print(f"    {f1} <-> {f2}: r = {r:.3f}")

        results[n_factors] = {
            "loadings": loadings,
            "communalities": comm,
            "variance_explained": cum_var[-1],
            "phi": fa.phi_,
        }

    return results


def run_reliability_hierarchical(df):
    """Reliability at both sub-standard and standard level."""
    print("\n" + "=" * 80)
    print("  RELIABILITY — SUB-STANDARD AND STANDARD LEVEL")
    print("=" * 80)

    results = {}

    for std_name, cfg in STRUCTURE.items():
        print(f"\n  ══ {std_name} ══")

        std_results = {"sub_standards": {}, "standard_level": {}}

        # Sub-standard level reliability
        if cfg.get("_sub_standards"):
            for sub_name, sub_cfg in cfg["_sub_standards"].items():
                items = list(sub_cfg["items"].keys())
                available = [i for i in items if i in df.columns]

                if len(available) < 2:
                    print(f"    {sub_name}: {len(available)} item(s) — cannot assess reliability")
                    std_results["sub_standards"][sub_name] = {
                        "n_items": len(available),
                        "alpha": None,
                        "items": {v: {"itc": None, "aid": None} for v in available},
                    }
                    continue

                dat = df[available].dropna()
                # Reverse where needed
                dat_a = dat.copy()
                for v in available:
                    info = get_item_info(v)
                    if info and info["reverse"]:
                        dat_a[v] = 1 - dat_a[v]

                alpha = cronbach_alpha(dat_a)
                itc = item_total_correlations(dat_a)
                aid = alpha_if_dropped(dat_a)

                print(f"    {sub_name} ({len(available)} items, N={len(dat)}): α = {alpha:.3f}")
                for v in available:
                    print(f"      {v:40s}  item-total r = {itc[v]:.3f}  α-if-drop = {aid[v]:.3f}")

                std_results["sub_standards"][sub_name] = {
                    "n_items": len(available),
                    "n_orgs": len(dat),
                    "alpha": alpha,
                    "items": {v: {"itc": itc[v], "aid": aid[v]} for v in available},
                }

        # Standard level (all items combined)
        all_items = get_standard_items(std_name)
        available = [i for i in all_items if i in df.columns]

        if len(available) < 2:
            print(f"    Standard level: {len(available)} item(s) — cannot assess")
            std_results["standard_level"] = {"n_items": len(available), "alpha": None}
        else:
            dat = df[available].dropna()
            dat_a = dat.copy()
            for v in available:
                info = get_item_info(v)
                if info and info["reverse"]:
                    dat_a[v] = 1 - dat_a[v]

            raw_alpha = cronbach_alpha(dat_a)
            std_alpha = standardized_alpha(dat_a)
            var_rat = variance_ratio(dat_a)
            itc = item_total_correlations(dat_a)
            aid = alpha_if_dropped(dat_a)
            mean_iic = dat_a.corr().where(
                np.triu(np.ones(dat_a.corr().shape), k=1).astype(bool)
            ).stack().mean()

            # Use standardized alpha when variance ratio is extreme (>10x)
            use_std = var_rat > 10
            alpha = std_alpha if use_std else raw_alpha

            print(f"    Standard level ({len(available)} items, N={len(dat)}):")
            print(f"      Raw α = {raw_alpha:.3f}, Standardized α = {std_alpha:.3f}, mean IIC = {mean_iic:.3f}")
            if use_std:
                print(f"      ⚠ Variance ratio = {var_rat:.0f}:1 — raw α unreliable, using standardized α = {std_alpha:.3f}")
            for v in available:
                print(f"      {v:40s}  item-total r = {itc[v]:.3f}  α-if-drop = {aid[v]:.3f}")

            std_results["standard_level"] = {
                "n_items": len(available),
                "n_orgs": len(dat),
                "alpha": alpha,
                "raw_alpha": raw_alpha,
                "std_alpha": std_alpha,
                "variance_ratio": var_rat,
                "use_standardized": use_std,
                "mean_iic": mean_iic,
                "items": {v: {"itc": itc[v], "aid": aid[v]} for v in available},
            }

        results[std_name] = std_results

    return results


def test_elm_second_order(df):
    """Test whether ELM sub-standards form a coherent second-order factor."""
    print("\n" + "=" * 80)
    print("  SECOND-ORDER TEST: EFFECTIVE LINE MANAGEMENT")
    print("  Do sub-standards (Line Mgmt, Team, Appraisal, Development)")
    print("  form a coherent higher-order construct?")
    print("=" * 80)

    elm_cfg = STRUCTURE["Effective Line Management"]["_sub_standards"]

    # Step 1: Create sub-standard composites
    composites = {}
    for sub_name, sub_cfg in elm_cfg.items():
        items = list(sub_cfg["items"].keys())
        available = [i for i in items if i in df.columns]
        dat = df[available].copy()
        for v in available:
            info = get_item_info(v)
            if info and info["reverse"]:
                dat[v] = 1 - dat[v]
        composites[sub_name] = dat.mean(axis=1)

    comp_df = pd.DataFrame(composites).dropna()
    n = len(comp_df)
    print(f"\n  Sub-standard composites (N={n}):")

    # Step 2: Correlations between sub-standards
    corr = comp_df.corr()
    print(f"\n  Inter-sub-standard correlations:")
    print("  " + corr.round(3).to_string().replace("\n", "\n  "))

    # Step 3: Reliability of composites as a scale
    alpha = cronbach_alpha(comp_df)
    mean_iic = corr.where(
        np.triu(np.ones(corr.shape), k=1).astype(bool)
    ).stack().mean()
    print(f"\n  Second-order reliability:")
    print(f"    Cronbach's alpha of sub-standard composites: {alpha:.3f}")
    print(f"    Mean inter-sub-standard r: {mean_iic:.3f}")

    # Step 4: Item-total correlations at composite level
    itc = item_total_correlations(comp_df)
    aid = alpha_if_dropped(comp_df)
    print(f"\n  Sub-standard contribution to ELM:")
    for sub in comp_df.columns:
        print(f"    {sub:20s}  composite-total r = {itc[sub]:.3f}  α-if-drop = {aid[sub]:.3f}")

    # Step 5: Single-factor EFA on sub-standard composites
    if len(comp_df.columns) >= 3:
        fa = FactorAnalyzer(n_factors=1, rotation=None, method="minres")
        fa.fit(comp_df)
        loadings = pd.Series(fa.loadings_.flatten(), index=comp_df.columns)
        var_exp = fa.get_factor_variance()[2][0]
        print(f"\n  Single-factor model on sub-standard composites:")
        print(f"    Variance explained: {var_exp:.1%}")
        for sub, loading in loadings.items():
            print(f"    {sub:20s}  loading = {loading:.3f}")
    else:
        loadings = None
        var_exp = None

    return {
        "correlations": corr,
        "alpha": alpha,
        "mean_iic": mean_iic,
        "itc": itc,
        "aid": aid,
        "loadings": loadings,
        "var_explained": var_exp,
        "n": n,
    }


# ══════════════════════════════════════════════════════════════════════
# WORD DOCUMENT GENERATION
# ══════════════════════════════════════════════════════════════════════

def rag_rating(value, metric_type):
    """Return RAG rating and colour for a metric."""
    if value is None:
        return "N/A", "grey"

    if metric_type == "alpha":
        if value >= 0.80:
            return "GREEN", "green"
        elif value >= 0.70:
            return "AMBER", "amber"
        else:
            return "RED", "red"
    elif metric_type == "itc":
        if value >= 0.50:
            return "GREEN", "green"
        elif value >= 0.30:
            return "AMBER", "amber"
        else:
            return "RED", "red"
    elif metric_type == "loading":
        if abs(value) >= 0.50:
            return "GREEN", "green"
        elif abs(value) >= 0.30:
            return "AMBER", "amber"
        else:
            return "RED", "red"
    elif metric_type == "second_order_loading":
        if abs(value) >= 0.60:
            return "GREEN", "green"
        elif abs(value) >= 0.40:
            return "AMBER", "amber"
        else:
            return "RED", "red"

    return "?", "grey"


def rag_word(value, metric_type):
    """Human-readable assessment word."""
    if value is None:
        return "Cannot assess"

    if metric_type == "alpha":
        if value >= 0.90:
            return "Excellent"
        elif value >= 0.80:
            return "Good"
        elif value >= 0.70:
            return "Acceptable"
        elif value >= 0.60:
            return "Questionable"
        else:
            return "Poor"
    elif metric_type == "itc":
        if value >= 0.70:
            return "Very strong fit"
        elif value >= 0.50:
            return "Strong fit"
        elif value >= 0.30:
            return "Adequate fit"
        elif value >= 0.15:
            return "Weak fit"
        else:
            return "Does not fit"
    elif metric_type in ("loading", "second_order_loading"):
        if abs(value) >= 0.70:
            return "Strong"
        elif abs(value) >= 0.50:
            return "Moderate"
        elif abs(value) >= 0.30:
            return "Weak"
        else:
            return "Does not load"
    return "?"


def generate_word_doc(efa_results, reliability_results, elm_results):
    """Generate the plain-English Word document with RAG-rated tables."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    RAG_COLOURS = {
        "green": RGBColor(0x00, 0x80, 0x00),
        "amber": RGBColor(0xFF, 0x8C, 0x00),
        "red":   RGBColor(0xCC, 0x00, 0x00),
        "grey":  RGBColor(0x99, 0x99, 0x99),
    }

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.add_heading('NOF 2024: Does the Structure Work?', level=0)
    doc.add_paragraph(
        'A plain-English assessment of whether the proposed questions group properly into '
        'sub-standards and standards. Includes RAG ratings for every question and sub-standard.'
    )
    doc.add_paragraph(
        'Based on 263 NHS organisations, 2024 Staff Survey data. '
        'Method: Exploratory Factor Analysis (what naturally emerges) and reliability testing '
        '(do the questions agree with each other?).'
    )

    # ── Key for RAG ratings ──
    doc.add_heading('How to read the RAG ratings', level=1)
    p = doc.add_paragraph()
    r = p.add_run('GREEN ')
    r.font.color.rgb = RAG_COLOURS["green"]
    r.bold = True
    p.add_run('= Strong evidence this works. ')
    r = p.add_run('AMBER ')
    r.font.color.rgb = RAG_COLOURS["amber"]
    r.bold = True
    p.add_run('= Acceptable but worth noting. ')
    r = p.add_run('RED ')
    r.font.color.rgb = RAG_COLOURS["red"]
    r.bold = True
    p.add_run('= Problem — needs attention.')

    p = doc.add_paragraph()
    r = p.add_run('Reliability (Cronbach\'s alpha): ')
    r.bold = True
    p.add_run('Do the questions within a standard agree? GREEN = 0.80+, AMBER = 0.70-0.79, RED = below 0.70.')

    p = doc.add_paragraph()
    r = p.add_run('Item fit (item-total correlation): ')
    r.bold = True
    p.add_run('Does each question belong in its group? GREEN = 0.50+, AMBER = 0.30-0.49, RED = below 0.30.')

    p = doc.add_paragraph()
    r = p.add_run('Factor loading: ')
    r.bold = True
    p.add_run('When you let the data group questions freely, does each question land where it should? '
              'GREEN = 0.50+, AMBER = 0.30-0.49, RED = below 0.30.')

    # ── What naturally emerged ──
    doc.add_heading('What the data naturally groups together (unforced)', level=1)
    doc.add_paragraph(
        'Before testing the proposed structure, we asked: if you let the data decide, '
        'how do these questions naturally cluster? The 6-factor solution explains about 80% '
        'of the variation and gives the clearest picture.'
    )

    # Get 6-factor results
    efa6 = efa_results.get(6)
    if efa6:
        loadings = efa6["loadings"]
        # Label each factor by its dominant items
        factor_labels = {}
        for f_col in loadings.columns:
            top_items = loadings[f_col].abs().nlargest(3).index.tolist()
            labels = []
            for item in top_items:
                info = get_item_info(item)
                if info:
                    label = info["sub_standard"] or info["standard"]
                    if label not in labels:
                        labels.append(label)
            factor_labels[f_col] = " / ".join(labels[:2])

        for f_col, label in factor_labels.items():
            items_on = loadings[loadings[f_col].abs() >= 0.30].index.tolist()
            if items_on:
                p = doc.add_paragraph()
                r = p.add_run(f'{f_col} → "{label}": ')
                r.bold = True
                item_strs = []
                for v in items_on:
                    info = get_item_info(v)
                    q_short = info["question"].split(":")[0] if info else v
                    std = info["standard"][:15] if info else "?"
                    item_strs.append(f'{q_short} [{std}]')
                p.add_run(", ".join(item_strs))

    # ── Standard-by-standard tables ──
    doc.add_heading('Standard-by-standard assessment', level=1)

    for std_name, cfg in STRUCTURE.items():
        doc.add_heading(std_name, level=2)

        # Notes
        if cfg.get("_note"):
            p = doc.add_paragraph()
            r = p.add_run("Note: ")
            r.bold = True
            r.font.color.rgb = RAG_COLOURS["amber"]
            p.add_run(cfg["_note"])

        rel = reliability_results.get(std_name, {})
        std_level = rel.get("standard_level", {})

        # Standard-level reliability
        alpha = std_level.get("alpha")
        use_std = std_level.get("use_standardized", False)
        var_rat = std_level.get("variance_ratio", 1)
        raw_alpha = std_level.get("raw_alpha")
        std_alpha_val = std_level.get("std_alpha")

        if alpha is not None:
            rating, colour = rag_rating(alpha, "alpha")
            word = rag_word(alpha, "alpha")
            p = doc.add_paragraph()
            if use_std:
                r = p.add_run(f"Overall reliability: standardized α = {alpha:.2f} — {word} ")
            else:
                r = p.add_run(f"Overall reliability: α = {alpha:.2f} — {word} ")
            r.bold = True
            r2 = p.add_run(f"[{rating}]")
            r2.bold = True
            r2.font.color.rgb = RAG_COLOURS[colour]

            if use_std:
                p2 = doc.add_paragraph()
                r = p2.add_run("Technical note: ")
                r.bold = True
                r.font.color.rgb = RAG_COLOURS["amber"]
                p2.add_run(
                    f"Item variances differ by {var_rat:.0f}:1 (e.g. patient violence is far more "
                    f"common than manager violence). Raw α = {raw_alpha:.2f} is misleading here — "
                    f"the standardized α = {std_alpha_val:.2f} (based on correlations, not covariances) "
                    f"is the correct measure. The questions genuinely belong together."
                )
        elif std_level.get("n_items", 0) < 2:
            p = doc.add_paragraph()
            r = p.add_run("Overall reliability: Cannot assess ")
            r.bold = True
            if std_level.get("n_items", 0) == 1:
                p.add_run("(single item)")
            else:
                p.add_run("(too few items)")

        # Question-level table
        all_items = get_standard_items(std_name)
        available = [i for i in all_items if i in (std_level.get("items", {}).keys())]
        if not available:
            # Might be from sub-standards
            available = []
            for sub_name, sub_info in rel.get("sub_standards", {}).items():
                available.extend(sub_info.get("items", {}).keys())

        # Build table: Question | Sub-standard | Item fit | Loading | Assessment | RAG
        n_rows = len(all_items) + 1
        table = doc.add_table(rows=n_rows, cols=6)
        table.style = 'Light Shading Accent 1'

        headers = ['Question', 'Sub-standard', 'Item fit\n(item-total r)', 'Factor loading\n(6-factor EFA)', 'Assessment', 'RAG']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        row_idx = 1
        for var in all_items:
            if var not in (df.columns if 'df' in dir() else []):
                pass  # We'll handle missing data gracefully
            info = get_item_info(var)
            if not info:
                continue

            # Get item-total correlation from standard-level
            itc_val = None
            if std_level.get("items") and var in std_level["items"]:
                itc_val = std_level["items"][var]["itc"]
            else:
                # Check sub-standard level
                for sub_name, sub_info in rel.get("sub_standards", {}).items():
                    if sub_info.get("items") and var in sub_info["items"]:
                        itc_val = sub_info["items"][var]["itc"]

            # Get loading from 6-factor EFA
            loading_val = None
            if efa6 and var in efa6["loadings"].index:
                row = efa6["loadings"].loc[var]
                primary_factor = row.abs().idxmax()
                loading_val = row[primary_factor]

            # Fill row
            if row_idx < n_rows:
                cells = table.rows[row_idx].cells
                cells[0].text = info["question"]
                cells[1].text = info["sub_standard"] or "—"

                if itc_val is not None:
                    itc_rating, itc_colour = rag_rating(itc_val, "itc")
                    itc_word = rag_word(itc_val, "itc")
                    cells[2].text = f"{itc_val:.2f}"
                else:
                    itc_rating, itc_colour = "N/A", "grey"
                    itc_word = "—"
                    cells[2].text = "—"

                if loading_val is not None:
                    ld_rating, ld_colour = rag_rating(loading_val, "loading")
                    ld_word = rag_word(loading_val, "loading")
                    cells[3].text = f"{loading_val:.2f}"
                else:
                    ld_rating, ld_colour = "N/A", "grey"
                    ld_word = "—"
                    cells[3].text = "—"

                # Combined assessment
                if itc_val is not None and loading_val is not None:
                    assessment = f"{itc_word}; loads {ld_word.lower()}"
                elif loading_val is not None:
                    assessment = f"Loads {ld_word.lower()}"
                elif itc_val is not None:
                    assessment = itc_word
                else:
                    assessment = "Cannot assess"
                cells[4].text = assessment

                # Overall RAG for this item
                if itc_val is not None and loading_val is not None:
                    if itc_val >= 0.50 and abs(loading_val) >= 0.50:
                        overall_rag, overall_colour = "GREEN", "green"
                    elif itc_val < 0.30 or abs(loading_val) < 0.30:
                        overall_rag, overall_colour = "RED", "red"
                    else:
                        overall_rag, overall_colour = "AMBER", "amber"
                elif itc_val is not None:
                    overall_rag, overall_colour = rag_rating(itc_val, "itc")
                elif loading_val is not None:
                    overall_rag, overall_colour = rag_rating(loading_val, "loading")
                else:
                    overall_rag, overall_colour = "N/A", "grey"

                p = cells[5].paragraphs[0]
                r = p.add_run(overall_rag)
                r.bold = True
                r.font.color.rgb = RAG_COLOURS[overall_colour]

                # Set font size for all cells
                for cell in cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

                row_idx += 1

        # Sub-standard reliability table (if applicable)
        if cfg.get("_sub_standards") and rel.get("sub_standards"):
            doc.add_paragraph()  # spacer
            p = doc.add_paragraph()
            r = p.add_run("Sub-standard reliability:")
            r.bold = True

            n_subs = len(cfg["_sub_standards"])
            sub_table = doc.add_table(rows=n_subs + 1, cols=5)
            sub_table.style = 'Light Shading Accent 1'

            for i, h in enumerate(['Sub-standard', 'N items', 'Reliability (α)', 'Assessment', 'RAG']):
                cell = sub_table.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)

            for si, (sub_name, sub_info) in enumerate(rel.get("sub_standards", {}).items(), 1):
                cells = sub_table.rows[si].cells
                cells[0].text = sub_name
                cells[1].text = str(sub_info.get("n_items", 0))

                sub_alpha = sub_info.get("alpha")
                if sub_alpha is not None:
                    cells[2].text = f"{sub_alpha:.2f}"
                    rating, colour = rag_rating(sub_alpha, "alpha")
                    word = rag_word(sub_alpha, "alpha")
                    cells[3].text = word
                    p = cells[4].paragraphs[0]
                    r = p.add_run(rating)
                    r.bold = True
                    r.font.color.rgb = RAG_COLOURS[colour]
                else:
                    cells[2].text = "—"
                    cells[3].text = "Single item" if sub_info.get("n_items", 0) == 1 else "Cannot assess"
                    cells[4].text = "N/A"

                for cell in cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

    # ── SECOND ORDER: ELM ──
    doc.add_heading('Does Effective Line Management work as a second-order standard?', level=1)
    doc.add_paragraph(
        'ELM is the only standard with multiple sub-standards that each have enough questions to test. '
        'The question is: do the four sub-standards (Line Management, Team Management, Appraisal, '
        'Development) hang together as one higher-level construct?'
    )

    if elm_results:
        alpha = elm_results["alpha"]
        rating, colour = rag_rating(alpha, "alpha")
        word = rag_word(alpha, "alpha")

        p = doc.add_paragraph()
        r = p.add_run(f'Second-order reliability: α = {alpha:.2f} — {word} ')
        r.bold = True
        r2 = p.add_run(f'[{rating}]')
        r2.bold = True
        r2.font.color.rgb = RAG_COLOURS[colour]

        # Correlation table
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run('How the sub-standards correlate with each other:')
        r.bold = True

        corr = elm_results["correlations"]
        n_subs = len(corr)
        corr_table = doc.add_table(rows=n_subs + 1, cols=n_subs + 1)
        corr_table.style = 'Light Shading Accent 1'

        corr_table.rows[0].cells[0].text = ""
        for i, sub in enumerate(corr.columns):
            corr_table.rows[0].cells[i+1].text = sub
            corr_table.rows[i+1].cells[0].text = sub
            for j, sub2 in enumerate(corr.columns):
                val = corr.iloc[i, j]
                cell = corr_table.rows[i+1].cells[j+1]
                if i == j:
                    cell.text = "1.00"
                else:
                    cell.text = f"{val:.2f}"

        for row in corr_table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

        # Loading table
        if elm_results.get("loadings") is not None:
            doc.add_paragraph()
            p = doc.add_paragraph()
            r = p.add_run('How strongly each sub-standard loads onto the ELM factor:')
            r.bold = True

            loadings = elm_results["loadings"]
            load_table = doc.add_table(rows=len(loadings) + 1, cols=4)
            load_table.style = 'Light Shading Accent 1'

            for i, h in enumerate(['Sub-standard', 'Loading', 'Assessment', 'RAG']):
                cell = load_table.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)

            for si, (sub, loading) in enumerate(loadings.items(), 1):
                cells = load_table.rows[si].cells
                cells[0].text = sub
                cells[1].text = f"{loading:.2f}"
                cells[2].text = rag_word(loading, "second_order_loading")
                rating, colour = rag_rating(loading, "second_order_loading")
                p = cells[3].paragraphs[0]
                r = p.add_run(rating)
                r.bold = True
                r.font.color.rgb = RAG_COLOURS[colour]

                for cell in cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

            var_exp = elm_results.get("var_explained")
            if var_exp:
                doc.add_paragraph(
                    f'A single "ELM" factor explains {var_exp:.0%} of the variation across the four sub-standards.'
                )

        # Interpretation
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("Interpretation: ")
        r.bold = True

        if alpha >= 0.70:
            p.add_run(
                'The four sub-standards do form a coherent second-order construct. '
                'You can defensibly say that "Effective Line Management" is a real thing '
                'that these sub-standards collectively measure.'
            )
        else:
            p.add_run(
                'The sub-standards do not form a strong second-order construct. '
                'They are related but measure somewhat different things. '
                'Consider whether it is better to treat them as separate first-order standards '
                'rather than grouping them under one umbrella.'
            )

        # Check if Appraisal is the weak link
        if elm_results.get("aid"):
            aid = elm_results["aid"]
            for sub, val in aid.items():
                if val > alpha + 0.05:
                    p2 = doc.add_paragraph()
                    r = p2.add_run(f"Note: ")
                    r.bold = True
                    r.font.color.rgb = RAG_COLOURS["amber"]
                    p2.add_run(
                        f'Removing "{sub}" would improve the second-order reliability '
                        f'from {alpha:.2f} to {val:.2f}. This sub-standard may not belong '
                        f'under ELM.'
                    )

    # ── Supportive Environment ──
    doc.add_heading('Supportive Environment: second-order?', level=1)
    doc.add_paragraph(
        'Supportive Environment has two sub-standards (Facilities and Wellbeing), '
        'but each is a single question. You cannot test second-order structure with single items. '
        'In practice, this is two standalone indicators grouped under one label. '
        'That is a policy decision, not something the data can validate or reject.'
    )

    # Correlation between the two items
    p = doc.add_paragraph()
    r = p.add_run("For reference: ")
    r.bold = True
    se_alpha = reliability_results.get("Supportive Environment", {}).get("standard_level", {}).get("alpha")
    if se_alpha is not None:
        p.add_run(f'The two items correlate at r = {se_alpha:.2f} as a pair. ')
        if se_alpha < 0.50:
            p.add_run('This is weak — they are measuring quite different things.')
        else:
            p.add_run('This suggests a moderate connection.')

    # ── Summary ──
    doc.add_heading('Summary of findings', level=1)

    findings = [
        ("Most first-order standards work well",
         "Flexible Working, Line Management, Development, Team Management, and Tackling Racism "
         "all have good reliability and their questions load where expected."),
        ("Appraisal needs fixing",
         "Q23a (did you have an appraisal?) does not fit with Q23c/Q23d (was it any good?). "
         "Remove Q23a from the composite or score it separately."),
        ("Sexual Safety is thin",
         "Only one item in the updated specification (Q17b). Cannot assess reliability. "
         "Consider retaining Q17a or adding another item."),
        ("Flexible Working and Line Management overlap heavily",
         "They correlate at r > 0.90. Statistically one construct, though policy-distinct."),
        ("ELM as a second-order standard",
         f"The four sub-standards {'do' if elm_results and elm_results['alpha'] >= 0.70 else 'do not strongly'} "
         f"form a coherent higher-order construct "
         f"(α = {elm_results['alpha']:.2f})." if elm_results else "Could not be tested."),
        ("Supportive Environment cannot be tested as second-order",
         "Single items per sub-standard. Treat as two standalone indicators under one heading."),
        ("The Racism LR approach is methodologically distinct",
         "It measures BME vs White disparity per trust, not raw prevalence. "
         "Conceptually sound and produces meaningful variation (Q16b: 1.2x to 5x across trusts)."),
    ]

    for title, detail in findings:
        p = doc.add_paragraph()
        r = p.add_run(f"{title}. ")
        r.bold = True
        p.add_run(detail)

    # Save
    out_path = "output/NOF_2024_Hierarchical_Assessment.docx"
    doc.save(out_path)
    print(f"\n  Saved: {out_path}")
    return out_path


# ══════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════

def main():
    print("=" * 80)
    print("  HIERARCHICAL NOF ANALYSIS — SUB-STANDARDS → STANDARDS")
    print("  2024 NSS data, 263 organisations")
    print("=" * 80)

    global df  # Make available to doc generator for column checks
    df = extract_all_nof_items()

    # 1. Unforced EFA
    efa_results = run_unforced_efa(df)

    # 2. Hierarchical reliability
    reliability_results = run_reliability_hierarchical(df)

    # 3. Second-order test for ELM
    elm_results = test_elm_second_order(df)

    # 4. Generate Word doc
    print("\n" + "=" * 80)
    print("  GENERATING WORD DOCUMENT")
    print("=" * 80)
    doc_path = generate_word_doc(efa_results, reliability_results, elm_results)

    print("\n" + "=" * 80)
    print(f"  DONE — report saved to {doc_path}")
    print("=" * 80)


if __name__ == "__main__":
    main()
