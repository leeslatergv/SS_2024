"""
Generate formatted outputs from the 2024 NOF analysis.
Produces:
  1. Excel workbook with multiple tabs (all data tables)
  2. Word report summarising findings
  3. CSVs for individual tables

Re-uses the extraction and scoring from nof_2024_only.py.
"""

import pandas as pd
import numpy as np
from factor_analyzer import FactorAnalyzer
from factor_analyzer.factor_analyzer import calculate_kmo
from scipy import stats
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import xlsxwriter
import os
import warnings
warnings.filterwarnings("ignore")

# Import extraction functions from main script
from nof_2024_only import (
    extract_all_nof_items, cronbach_alpha, item_total_correlations,
    alpha_if_dropped
)

OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ── Standard definitions ────────────────────────────────────────────
STANDARDS = {
    "Reducing Violence": {
        "items": ["violence_patients", "violence_managers", "violence_colleagues",
                  "harassment_patients", "harassment_managers", "harassment_colleagues"],
        "reverse": ["violence_patients", "violence_managers", "violence_colleagues",
                    "harassment_patients", "harassment_managers", "harassment_colleagues"],
        "questions": {
            "violence_patients": "Q13a: Physical violence from patients/public",
            "violence_managers": "Q13b: Physical violence from managers",
            "violence_colleagues": "Q13c: Physical violence from colleagues",
            "harassment_patients": "Q14a: Harassment from patients/public",
            "harassment_managers": "Q14b: Harassment from managers",
            "harassment_colleagues": "Q14c: Harassment from colleagues",
        },
        "extra_items": {
            "violence_reported": "Q13d: Violence reported",
        },
    },
    "Tackling Racism": {
        "items": ["fair_career_progression", "discrim_patients",
                  "discrim_colleagues", "org_respects_differences"],
        "reverse": ["discrim_patients", "discrim_colleagues"],
        "questions": {
            "fair_career_progression": "Q15: Org acts fairly re career progression",
            "discrim_patients": "Q16a: Discrimination from patients/public",
            "discrim_colleagues": "Q16b: Discrimination from manager/colleagues",
            "org_respects_differences": "Q21: Org respects individual differences",
        },
    },
    "Sexual Safety": {
        "items": ["sexual_behaviour_patients", "sexual_behaviour_staff"],
        "reverse": ["sexual_behaviour_patients", "sexual_behaviour_staff"],
        "questions": {
            "sexual_behaviour_patients": "Q17a: Unwanted sexual behaviour from patients",
            "sexual_behaviour_staff": "Q17b: Unwanted sexual behaviour from staff",
        },
    },
    "Flexible Working": {
        "items": ["satisfaction_flexible_working", "org_committed_wlb",
                  "manager_flexible_working"],
        "reverse": [],
        "questions": {
            "satisfaction_flexible_working": "Q4d: Satisfaction with flexible working",
            "org_committed_wlb": "Q6b: Org committed to helping balance work/home",
            "manager_flexible_working": "Q6d: Can approach manager about flexible working",
        },
        "comparison_item": {
            "good_wl_balance": "Q6c: Achieve good work-life balance",
        },
    },
    "Line Management": {
        "items": ["mgr_interest_wellbeing", "mgr_understanding_problems",
                  "mgr_effective_action"],
        "reverse": [],
        "questions": {
            "mgr_interest_wellbeing": "Q9d: Manager takes interest in wellbeing",
            "mgr_understanding_problems": "Q9f: Manager works to understand problems",
            "mgr_effective_action": "Q9i: Manager takes effective action",
        },
    },
    "Team Management": {
        "items": ["team_shared_objectives", "feel_valued_by_team"],
        "reverse": [],
        "questions": {
            "team_shared_objectives": "Q7a: Team has shared objectives",
            "feel_valued_by_team": "Q7h: I feel valued by my team",
        },
    },
    "Appraisal": {
        "items": ["had_appraisal", "appraisal_clear_objectives", "appraisal_felt_valued"],
        "reverse": [],
        "questions": {
            "had_appraisal": "Q23a: Had appraisal in last 12 months",
            "appraisal_clear_objectives": "Q23c: Appraisal helped agree clear objectives",
            "appraisal_felt_valued": "Q23d: Appraisal left feeling valued",
        },
    },
    "Development": {
        "items": ["supported_develop_potential", "access_learning_dev"],
        "reverse": [],
        "questions": {
            "supported_develop_potential": "Q24d: Feel supported to develop potential",
            "access_learning_dev": "Q24e: Can access learning/development opportunities",
        },
    },
    "Supportive Env - Facilities": {
        "items": ["nutritious_food"],
        "reverse": [],
        "questions": {
            "nutritious_food": "Q22: Can eat nutritious/affordable food at work",
        },
    },
    "Supportive Env - Wellbeing": {
        "items": ["org_positive_wellbeing"],
        "reverse": [],
        "questions": {
            "org_positive_wellbeing": "Q11a: Org takes positive action on wellbeing",
        },
    },
}


def compute_all_results(df):
    """Compute all analysis results and return as dict of DataFrames."""
    results = {}

    # ── 1. Descriptive statistics ────────────────────────────────────
    all_items = []
    all_questions = {}
    all_standards = {}
    for std_name, cfg in STANDARDS.items():
        for item in cfg["items"]:
            if item not in all_items:
                all_items.append(item)
                all_questions[item] = cfg["questions"].get(item, item)
                all_standards[item] = std_name
        if "extra_items" in cfg:
            for item, q in cfg["extra_items"].items():
                if item not in all_items:
                    all_items.append(item)
                    all_questions[item] = q
                    all_standards[item] = std_name
        if "comparison_item" in cfg:
            for item, q in cfg["comparison_item"].items():
                if item not in all_items:
                    all_items.append(item)
                    all_questions[item] = q
                    all_standards[item] = std_name + " (comparison)"

    desc_rows = []
    for item in all_items:
        if item in df.columns:
            s = df[item].dropna()
            desc_rows.append({
                "Standard": all_standards.get(item, ""),
                "Variable": item,
                "Question": all_questions.get(item, ""),
                "N": len(s),
                "Mean": s.mean(),
                "SD": s.std(),
                "Min": s.min(),
                "Max": s.max(),
            })
    results["Descriptive Stats"] = pd.DataFrame(desc_rows)

    # ── 2. Reliability per standard ──────────────────────────────────
    rel_rows = []
    summary_rows = []
    for std_name, cfg in STANDARDS.items():
        items = cfg["items"]
        rev = cfg.get("reverse", [])
        available = [i for i in items if i in df.columns]
        if len(available) < 2:
            summary_rows.append({
                "Standard": std_name, "N Items": len(available),
                "N Orgs": df[available].dropna().shape[0] if available else 0,
                "Cronbach Alpha": np.nan, "Mean Inter-Item r": np.nan,
                "Verdict": "Single item — no reliability"
            })
            continue

        dat = df[available].dropna()
        dat_aligned = dat.copy()
        for ri in rev:
            if ri in dat_aligned.columns:
                dat_aligned[ri] = 1 - dat_aligned[ri]

        alpha = cronbach_alpha(dat_aligned)
        itc = item_total_correlations(dat_aligned)
        aid = alpha_if_dropped(dat_aligned)
        mean_iic = dat_aligned.corr().where(
            np.triu(np.ones(dat_aligned.corr().shape), k=1).astype(bool)
        ).stack().mean()

        if alpha >= 0.8:
            verdict = "Good"
        elif alpha >= 0.7:
            verdict = "Acceptable"
        elif alpha >= 0.6:
            verdict = "Questionable"
        elif alpha >= 0.5:
            verdict = "Poor"
        else:
            verdict = "Unacceptable"

        summary_rows.append({
            "Standard": std_name, "N Items": len(available),
            "N Orgs": len(dat), "Cronbach Alpha": alpha,
            "Mean Inter-Item r": mean_iic, "Verdict": verdict,
        })

        for item in available:
            rel_rows.append({
                "Standard": std_name,
                "Variable": item,
                "Question": cfg["questions"].get(item, item),
                "Reversed": "Yes" if item in rev else "No",
                "Item-Total r": itc[item],
                "Alpha if Dropped": aid[item],
                "Standard Alpha": alpha,
                "Fits Well": "Yes" if itc[item] > 0.3 else "NO — weak fit",
            })

    results["Reliability Summary"] = pd.DataFrame(summary_rows)
    results["Item Reliability Detail"] = pd.DataFrame(rel_rows)

    # ── 3. Q6b vs Q6c comparison ─────────────────────────────────────
    q6_rows = []
    for label, flex_item in [("With Q6c (good_wl_balance)", "good_wl_balance"),
                              ("With Q6b (org_committed_wlb)", "org_committed_wlb")]:
        items = ["satisfaction_flexible_working", flex_item, "manager_flexible_working"]
        available = [i for i in items if i in df.columns]
        dat = df[available].dropna()
        alpha = cronbach_alpha(dat)
        itc = item_total_correlations(dat)
        mean_iic = dat.corr().where(
            np.triu(np.ones(dat.corr().shape), k=1).astype(bool)
        ).stack().mean()
        for item in available:
            q6_rows.append({
                "Configuration": label, "Variable": item,
                "Item-Total r": itc[item], "Cronbach Alpha": alpha,
                "Mean Inter-Item r": mean_iic, "N": len(dat),
            })
    results["Q6b vs Q6c"] = pd.DataFrame(q6_rows)

    # Also the inter-correlations
    flex_items = ["satisfaction_flexible_working", "good_wl_balance",
                  "org_committed_wlb", "manager_flexible_working"]
    avail = [i for i in flex_items if i in df.columns]
    flex_corr = df[avail].dropna().corr()
    flex_corr.index.name = "Variable"
    results["Flex Item Correlations"] = flex_corr.reset_index()

    # ── 4. EFA loadings (6-factor) ───────────────────────────────────
    efa_items = [
        "violence_patients", "violence_managers", "violence_colleagues",
        "harassment_patients", "harassment_managers", "harassment_colleagues",
        "fair_career_progression", "discrim_patients", "discrim_colleagues",
        "org_respects_differences",
        "sexual_behaviour_patients", "sexual_behaviour_staff",
        "satisfaction_flexible_working", "org_committed_wlb", "manager_flexible_working",
        "mgr_interest_wellbeing", "mgr_understanding_problems", "mgr_effective_action",
        "team_shared_objectives", "feel_valued_by_team",
        "had_appraisal", "appraisal_clear_objectives", "appraisal_felt_valued",
        "supported_develop_potential", "access_learning_dev",
        "nutritious_food", "org_positive_wellbeing",
    ]
    available_efa = [i for i in efa_items if i in df.columns]
    dat_efa = df[available_efa].dropna()

    for n_factors, label in [(6, "6"), (8, "8")]:
        fa = FactorAnalyzer(n_factors=n_factors, rotation="promax", method="minres")
        fa.fit(dat_efa)

        loadings = pd.DataFrame(
            fa.loadings_, index=available_efa,
            columns=[f"Factor {i+1}" for i in range(n_factors)]
        )

        comm = pd.Series(fa.get_communalities(), index=available_efa)
        _, _, cum_var = fa.get_factor_variance()

        efa_rows = []
        for item in available_efa:
            row = loadings.loc[item]
            primary = row.abs().idxmax()
            primary_val = row[primary]
            std = all_standards.get(item, "")
            q = all_questions.get(item, item)

            entry = {
                "Standard (Theorised)": std,
                "Variable": item,
                "Question": q,
            }
            for col in loadings.columns:
                val = row[col]
                entry[col] = val if abs(val) >= 0.30 else np.nan
                entry[f"{col} (raw)"] = val
            entry["Primary Factor"] = primary
            entry["Primary Loading"] = primary_val
            entry["Communality"] = comm[item]
            entry["Heywood Case"] = "YES" if comm[item] > 1.0 else ""
            entry["Loads as Expected"] = "Yes" if abs(primary_val) >= 0.40 else "Weak"
            efa_rows.append(entry)

        results[f"EFA {label}-Factor Loadings"] = pd.DataFrame(efa_rows)

        # Factor correlations
        if fa.phi_ is not None:
            phi = pd.DataFrame(
                fa.phi_,
                index=[f"Factor {i+1}" for i in range(n_factors)],
                columns=[f"Factor {i+1}" for i in range(n_factors)]
            )
            phi.index.name = "Factor"
            results[f"EFA {label}-Factor Correlations"] = phi.reset_index()

    # ── 5. Inter-standard correlations ───────────────────────────────
    composites = {}
    for std_name, cfg in STANDARDS.items():
        available = [i for i in cfg["items"] if i in df.columns]
        if not available:
            continue
        vals = df[available].copy()
        for ri in cfg.get("reverse", []):
            if ri in vals.columns:
                vals[ri] = 1 - vals[ri]
        composites[std_name] = vals.mean(axis=1)

    scores = pd.DataFrame(composites)
    inter_corr = scores.dropna().corr()
    inter_corr.index.name = "Standard"
    results["Inter-Standard Correlations"] = inter_corr.reset_index()

    # ── 6. Full correlation matrix of all items ──────────────────────
    all_avail = [i for i in all_items if i in df.columns]
    full_corr = df[all_avail].corr()
    full_corr.index.name = "Variable"
    results["Full Item Correlations"] = full_corr.reset_index()

    # ── 7. Racism likelihood ratio (per-trust BME vs White) ────────
    try:
        BENCHMARK_FILE = "Benchmark-report-Excel-data-for-2020-2024 (3).xlsx"
        BENCHMARK_SHEETS = [
            "Acute&Acute Community Trusts",
            "Acute Specialist Trusts",
            "MH&LD, MH, LD & Community Trust",
            "Community Trusts",
            "Ambulance Trusts",
        ]

        # Column pairs: (White suffix _1_, BME suffix _2_) for 2024
        LR_INDICATORS = {
            "Q15: Fair career progression": {
                "white": "equal_eth_summary_1_2024",
                "bme": "equal_eth_summary_2_2024",
                "white_n": "equal_eth_summary_n_1_2024",
                "bme_n": "equal_eth_summary_n_2_2024",
                "direction": "positive",  # higher = better, so LR < 1 means BME disadvantage
            },
            "Q16b: Discrimination from manager/colleagues": {
                "white": "q16b_eth_summary_1_2024",
                "bme": "q16b_eth_summary_2_2024",
                "white_n": "q16b_eth_summary_n_1_2024",
                "bme_n": "q16b_eth_summary_n_2_2024",
                "direction": "negative",  # higher = worse, so LR > 1 means BME disadvantage
            },
            "Harassment from patients (by ethnicity)": {
                "white": "har_from_pat_eth_summary_1_2024",
                "bme": "har_from_pat_eth_summary_2_2024",
                "white_n": "har_from_pat_eth_summary_n_1_2024",
                "bme_n": "har_from_pat_eth_summary_n_2_2024",
                "direction": "negative",
            },
            "Harassment from staff (by ethnicity)": {
                "white": "har_from_staff_eth_summary_1_2024",
                "bme": "har_from_staff_eth_summary_2_2024",
                "white_n": "har_from_staff_eth_summary_n_1_2024",
                "bme_n": "har_from_staff_eth_summary_n_2_2024",
                "direction": "negative",
            },
        }

        # Read and concatenate all trust-type sheets
        all_benchmark = []
        for sheet in BENCHMARK_SHEETS:
            try:
                bm = pd.read_excel(BENCHMARK_FILE, sheet_name=sheet)
                bm['trust_type'] = sheet
                all_benchmark.append(bm)
            except Exception:
                pass

        if all_benchmark:
            benchmark = pd.concat(all_benchmark, ignore_index=True)

            # Identify org code column (first column)
            org_col = benchmark.columns[0]

            # Per-trust LR computation
            trust_lr_rows = []
            lr_summary_rows = []

            for indicator_name, cols in LR_INDICATORS.items():
                white_col = cols["white"]
                bme_col = cols["bme"]
                direction = cols["direction"]

                if white_col not in benchmark.columns or bme_col not in benchmark.columns:
                    continue

                white_vals = pd.to_numeric(benchmark[white_col], errors='coerce')
                bme_vals = pd.to_numeric(benchmark[bme_col], errors='coerce')

                # Get sample sizes if available
                white_n_col = cols.get("white_n", "")
                bme_n_col = cols.get("bme_n", "")
                has_n = white_n_col in benchmark.columns and bme_n_col in benchmark.columns

                valid = white_vals.notna() & bme_vals.notna() & (white_vals > 0)
                for idx in benchmark.index[valid]:
                    w = white_vals[idx]
                    b = bme_vals[idx]
                    lr = b / w if w > 0 else np.nan

                    row = {
                        'Indicator': indicator_name,
                        'Trust': benchmark.loc[idx, org_col],
                        'Trust Type': benchmark.loc[idx, 'trust_type'],
                        'White Rate': w,
                        'BME Rate': b,
                        'Likelihood Ratio (BME/White)': lr,
                    }
                    if has_n:
                        row['White N'] = pd.to_numeric(
                            benchmark.loc[idx, white_n_col], errors='coerce')
                        row['BME N'] = pd.to_numeric(
                            benchmark.loc[idx, bme_n_col], errors='coerce')
                    trust_lr_rows.append(row)

                # Summary statistics across trusts
                lr_vals = bme_vals[valid] / white_vals[valid]
                lr_vals = lr_vals.replace([np.inf, -np.inf], np.nan).dropna()

                if len(lr_vals) > 0:
                    if direction == "negative":
                        interp = (
                            f"BME staff {lr_vals.median():.1f}x more likely"
                            if lr_vals.median() > 1
                            else f"BME staff {1/lr_vals.median():.1f}x less likely"
                        )
                    else:
                        interp = (
                            f"BME staff {1/lr_vals.median():.1f}x less likely to report positively"
                            if lr_vals.median() < 1
                            else f"BME staff report more positively"
                        )

                    lr_summary_rows.append({
                        'Indicator': indicator_name,
                        'N Trusts': len(lr_vals),
                        'Median LR': lr_vals.median(),
                        'Mean LR': lr_vals.mean(),
                        'SD LR': lr_vals.std(),
                        'Min LR': lr_vals.min(),
                        'Q25 LR': lr_vals.quantile(0.25),
                        'Q75 LR': lr_vals.quantile(0.75),
                        'Max LR': lr_vals.max(),
                        'Direction': direction,
                        'Interpretation': interp,
                    })

            if lr_summary_rows:
                results["Racism LR Summary"] = pd.DataFrame(lr_summary_rows)

            if trust_lr_rows:
                results["Racism LR Per Trust"] = pd.DataFrame(trust_lr_rows)

    except Exception as e:
        print(f"  WARNING: Could not compute racism likelihood ratios: {e}")
        import traceback
        traceback.print_exc()

    # ── 8. KMO per item ─────────────────────────────────────────────
    try:
        kmo_per, kmo_total = calculate_kmo(dat_efa)
        kmo_rows = []
        for item, kmo_val in zip(available_efa, kmo_per):
            verdict = "Good" if kmo_val >= 0.7 else "Acceptable" if kmo_val >= 0.6 else "Mediocre" if kmo_val >= 0.5 else "Poor — consider removing"
            kmo_rows.append({
                "Variable": item,
                "Question": all_questions.get(item, item),
                "KMO": kmo_val,
                "Verdict": verdict,
            })
        kmo_rows.append({"Variable": "OVERALL", "Question": "", "KMO": kmo_total,
                         "Verdict": "Good" if kmo_total >= 0.7 else "Acceptable"})
        results["KMO Sampling Adequacy"] = pd.DataFrame(kmo_rows)
    except Exception:
        pass

    return results


def write_excel(results):
    """Write all results to a formatted Excel workbook."""
    path = os.path.join(OUTPUT_DIR, "NOF_2024_Analysis_Results.xlsx")
    writer = pd.ExcelWriter(path, engine='xlsxwriter')
    workbook = writer.book

    # Formats
    header_fmt = workbook.add_format({
        'bold': True, 'bg_color': '#003087', 'font_color': 'white',
        'border': 1, 'text_wrap': True, 'valign': 'top',
    })
    num_fmt = workbook.add_format({'num_format': '0.000', 'border': 1})
    int_fmt = workbook.add_format({'num_format': '0', 'border': 1})
    text_fmt = workbook.add_format({'border': 1, 'text_wrap': True, 'valign': 'top'})
    good_fmt = workbook.add_format({'bg_color': '#C6EFCE', 'border': 1, 'num_format': '0.000'})
    bad_fmt = workbook.add_format({'bg_color': '#FFC7CE', 'border': 1, 'num_format': '0.000'})
    warn_fmt = workbook.add_format({'bg_color': '#FFEB9C', 'border': 1, 'num_format': '0.000'})
    bold_fmt = workbook.add_format({'bold': True, 'border': 1, 'text_wrap': True})
    loading_fmt = workbook.add_format({'num_format': '0.000', 'border': 1, 'bg_color': '#D9E2F3'})

    # Tab order
    tab_order = [
        "Reliability Summary",
        "Item Reliability Detail",
        "Q6b vs Q6c",
        "Flex Item Correlations",
        "EFA 6-Factor Loadings",
        "EFA 6-Factor Correlations",
        "EFA 8-Factor Loadings",
        "EFA 8-Factor Correlations",
        "Inter-Standard Correlations",
        "Racism LR Summary",
        "Racism LR Per Trust",
        "Descriptive Stats",
        "KMO Sampling Adequacy",
        "Full Item Correlations",
    ]

    for sheet_name in tab_order:
        if sheet_name not in results:
            continue
        df = results[sheet_name]
        # Truncate sheet name to 31 chars (Excel limit)
        safe_name = sheet_name[:31]
        df.to_excel(writer, sheet_name=safe_name, index=False, startrow=1)

        ws = writer.sheets[safe_name]

        # Write headers
        for col_num, col_name in enumerate(df.columns):
            ws.write(0, col_num, col_name, header_fmt)

        # Set column widths
        for col_num, col_name in enumerate(df.columns):
            max_len = max(len(str(col_name)), df[col_name].astype(str).str.len().max())
            ws.set_column(col_num, col_num, min(max_len + 2, 45))

        # Conditional formatting for key sheets
        n_rows = len(df)
        if sheet_name == "Reliability Summary":
            # Color-code alpha
            alpha_col = df.columns.get_loc("Cronbach Alpha")
            for row in range(n_rows):
                val = df.iloc[row, alpha_col]
                if pd.notna(val):
                    fmt = good_fmt if val >= 0.7 else warn_fmt if val >= 0.5 else bad_fmt
                    ws.write(row + 1, alpha_col, val, fmt)

        elif sheet_name == "Item Reliability Detail":
            # Color-code item-total r and "Fits Well"
            itc_col = df.columns.get_loc("Item-Total r")
            fits_col = df.columns.get_loc("Fits Well")
            for row in range(n_rows):
                val = df.iloc[row, itc_col]
                if pd.notna(val):
                    fmt = good_fmt if val > 0.4 else warn_fmt if val > 0.2 else bad_fmt
                    ws.write(row + 1, itc_col, val, fmt)
                fits = df.iloc[row, fits_col]
                fmt = text_fmt if fits == "Yes" else bold_fmt
                ws.write(row + 1, fits_col, fits, fmt)

        elif "Loadings" in sheet_name:
            # Highlight significant loadings (>= 0.30)
            factor_cols = [c for c in df.columns if c.startswith("Factor") and "(raw)" not in c]
            for col_name in factor_cols:
                col_idx = df.columns.get_loc(col_name)
                for row in range(n_rows):
                    val = df.iloc[row, col_idx]
                    if pd.notna(val):
                        ws.write(row + 1, col_idx, val, loading_fmt)

        elif "Correlations" in sheet_name and "Item" not in sheet_name:
            # Color-code high correlations
            num_cols = df.select_dtypes(include=[np.number]).columns
            for col_name in num_cols:
                col_idx = df.columns.get_loc(col_name)
                for row in range(n_rows):
                    val = df.iloc[row, col_idx]
                    if pd.notna(val) and val != 1.0:
                        if abs(val) > 0.8:
                            ws.write(row + 1, col_idx, val, bad_fmt)
                        elif abs(val) > 0.6:
                            ws.write(row + 1, col_idx, val, warn_fmt)
                        else:
                            ws.write(row + 1, col_idx, val, num_fmt)

    writer.close()
    print(f"  Saved: {path}")
    return path


def write_csvs(results):
    """Write key tables as individual CSVs."""
    csv_dir = os.path.join(OUTPUT_DIR, "csv")
    os.makedirs(csv_dir, exist_ok=True)
    for name, df in results.items():
        safe = name.replace(" ", "_").replace("/", "_").replace("-", "_")
        path = os.path.join(csv_dir, f"{safe}.csv")
        df.to_csv(path, index=False)
    print(f"  Saved CSVs to: {csv_dir}/")


def add_table_to_doc(doc, df, title=None, small=False):
    """Add a DataFrame as a formatted table to a Word document."""
    if title:
        doc.add_heading(title, level=2)

    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    for j, col_name in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col_name)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8 if small else 9)

    # Data
    for i in range(len(df)):
        for j in range(len(df.columns)):
            cell = table.rows[i + 1].cells[j]
            val = df.iloc[i, j]
            if pd.isna(val):
                cell.text = ""
            elif isinstance(val, float):
                cell.text = f"{val:.3f}"
            else:
                cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(7 if small else 8)


def write_word(results):
    """Write a formatted Word report."""
    path = os.path.join(OUTPUT_DIR, "NOF_2024_Factor_Analysis_Report.docx")
    doc = Document()

    # Page setup - landscape for wide tables
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Title
    title = doc.add_heading('NHS Staff Standards: NOF Question Set Analysis', level=0)
    doc.add_paragraph('Data: NHS Staff Survey 2024 (organisational results)')
    doc.add_paragraph('Method: Exploratory Factor Analysis (minres extraction, promax rotation)')
    doc.add_paragraph(f'N = 263 organisations')
    doc.add_paragraph('')

    # ── Section 1: Reliability ────────────────────────────────────
    doc.add_heading('1. Reliability Analysis by Standard', level=1)
    doc.add_paragraph(
        "Cronbach's alpha measures internal consistency — whether items within each standard "
        "measure the same underlying construct. Values above 0.70 are generally considered "
        "acceptable; above 0.80 is good. Mean inter-item correlation (IIC) should be 0.15-0.50 "
        "for broad constructs."
    )

    if "Reliability Summary" in results:
        add_table_to_doc(doc, results["Reliability Summary"], "Summary")

    doc.add_paragraph('')
    doc.add_heading('Item-Level Reliability', level=2)
    doc.add_paragraph(
        "Item-total correlation shows how well each item correlates with the rest of its standard. "
        "Values below 0.30 suggest the item doesn't belong. 'Alpha if Dropped' shows what "
        "happens to the standard's reliability if that item is removed — if it goes UP, the item "
        "is weakening the scale."
    )

    if "Item Reliability Detail" in results:
        # Show subset of columns for readability
        detail = results["Item Reliability Detail"][
            ["Standard", "Variable", "Question", "Reversed", "Item-Total r",
             "Alpha if Dropped", "Fits Well"]
        ]
        add_table_to_doc(doc, detail, small=True)

    # ── Section 2: Q6b vs Q6c ────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('2. Q6b vs Q6c: Which Fits Better in Flexible Working?', level=1)
    doc.add_paragraph(
        "Q6b ('My organisation is committed to helping me balance my work and home life') "
        "was chosen by policy. Q6c ('I achieve a good balance between my work life and my "
        "home life') was used in the original paper. This section compares both."
    )

    if "Q6b vs Q6c" in results:
        add_table_to_doc(doc, results["Q6b vs Q6c"])

    doc.add_paragraph('')
    doc.add_paragraph(
        "Conclusion: Q6b produces higher Cronbach's alpha (0.960 vs 0.941) and better "
        "item-total correlations. Q6b and Q6c correlate at r = 0.971, so they measure "
        "almost the same thing, but Q6b coheres better with Q4d and Q6d because all three "
        "ask about organisational/managerial commitment rather than personal experience."
    )

    if "Flex Item Correlations" in results:
        add_table_to_doc(doc, results["Flex Item Correlations"],
                         "Inter-correlations Among Flexible Working Items")

    # ── Section 3: EFA ────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('3. Exploratory Factor Analysis', level=1)
    doc.add_paragraph(
        "EFA identifies clusters of items that co-vary together. The 'loading' shows how "
        "strongly each item belongs to each factor. Loadings >= 0.30 are shown; blanks mean "
        "the loading was below this threshold. The 'Primary Factor' column shows which factor "
        "each item loaded most strongly on."
    )
    doc.add_paragraph(
        "Items that load on a DIFFERENT factor than their theorised standard suggest the "
        "standard structure doesn't match how trusts actually vary on these measures."
    )

    if "EFA 6-Factor Loadings" in results:
        # Show the key columns
        efa6 = results["EFA 6-Factor Loadings"]
        display_cols = ["Standard (Theorised)", "Variable", "Question"]
        factor_cols = [c for c in efa6.columns if c.startswith("Factor") and "(raw)" not in c]
        display_cols += factor_cols
        display_cols += ["Primary Factor", "Primary Loading", "Communality", "Heywood Case"]
        efa6_display = efa6[display_cols].copy()
        add_table_to_doc(doc, efa6_display, "6-Factor Solution (promax rotation)", small=True)

    if "EFA 6-Factor Correlations" in results:
        doc.add_paragraph('')
        add_table_to_doc(doc, results["EFA 6-Factor Correlations"],
                         "6-Factor Correlations")
        doc.add_paragraph(
            "Factor correlations above 0.80 suggest two factors may be measuring the same "
            "thing and could be merged."
        )

    if "EFA 8-Factor Loadings" in results:
        doc.add_page_break()
        efa8 = results["EFA 8-Factor Loadings"]
        display_cols8 = ["Standard (Theorised)", "Variable"]
        factor_cols8 = [c for c in efa8.columns if c.startswith("Factor") and "(raw)" not in c]
        display_cols8 += factor_cols8
        display_cols8 += ["Primary Factor", "Primary Loading"]
        efa8_display = efa8[display_cols8].copy()
        add_table_to_doc(doc, efa8_display, "8-Factor Solution (promax rotation)", small=True)

    if "EFA 8-Factor Correlations" in results:
        doc.add_paragraph('')
        add_table_to_doc(doc, results["EFA 8-Factor Correlations"],
                         "8-Factor Correlations")

    # ── Section 4: Discriminant validity ──────────────────────────
    doc.add_page_break()
    doc.add_heading('4. Inter-Standard Correlations (Discriminant Validity)', level=1)
    doc.add_paragraph(
        "This shows how strongly each standard's composite score correlates with every other "
        "standard. High correlations (r > 0.80) suggest two standards may not be measuring "
        "distinct things at the trust level. Correlations above 0.80 are highlighted."
    )

    if "Inter-Standard Correlations" in results:
        add_table_to_doc(doc, results["Inter-Standard Correlations"])

    # ── Section 5: Racism Likelihood Ratio ──────────────────────────
    doc.add_page_break()
    doc.add_heading('5. Tackling Racism: Likelihood Ratio (BME vs White)', level=1)
    doc.add_paragraph(
        "The policy specifies that Tackling Racism items should use a BME vs White "
        "comparator — specifically a likelihood ratio: P(outcome | BME) / P(outcome | White). "
        "Using the NHS Staff Survey benchmark data, we computed per-trust likelihood ratios "
        "for each racism-related indicator where ethnicity breakdowns are available."
    )
    doc.add_paragraph(
        "A likelihood ratio (LR) > 1 for negative outcomes (discrimination, harassment) means "
        "BME staff are more likely to experience those outcomes. For positive outcomes (fair "
        "career progression), an LR < 1 means BME staff are less likely to report positively."
    )

    if "Racism LR Summary" in results:
        add_table_to_doc(doc, results["Racism LR Summary"],
                         "Likelihood Ratio Summary Across Trusts")

    doc.add_paragraph('')
    doc.add_paragraph(
        "The per-trust detail is provided in the Excel workbook ('Racism LR Per Trust' tab). "
        "Trusts with very small BME sample sizes may have unstable likelihood ratios."
    )
    doc.add_paragraph(
        "The factor analysis supports Q15, Q16a, Q16b, and Q21 belonging together as a "
        "Tackling Racism standard regardless of whether absolute rates or likelihood ratios "
        "are used — the underlying construct is the same, only the scoring metric differs."
    )

    # ── Section 6: Issues and recommendations ─────────────────────
    doc.add_page_break()
    doc.add_heading('6. Issues and Recommendations', level=1)

    issues = [
        ("Flexible Working + Line Management (r = 0.918)",
         "These two standards are almost perfectly correlated at trust level. Trusts that "
         "score well on flexible working almost always score well on line management. "
         "They function as a single 'good management' construct statistically, though they "
         "remain policy-distinct."),
        ("Reducing Violence (alpha = -0.060)",
         "The violence items (Q13a/b/c) have very different base rates — violence from patients "
         "is much more common than from managers/colleagues. They don't form a reliable scale. "
         "Consider: (a) using Q13a only, or (b) treating these as separate indicators rather "
         "than a composite."),
        ("Appraisal (alpha = 0.435)",
         "'Had appraisal' (Q23a, binary yes/no) doesn't cohere with the appraisal quality "
         "items (Q23c, Q23d). Alpha without Q23a rises to 0.914. Consider: (a) scoring Q23a "
         "separately, or (b) only including Q23c and Q23d in the composite."),
        ("Q6b vs Q6c",
         "Q6b (org committed to WLB) fits the Flexible Working standard slightly better than "
         "Q6c (achieve good balance). Both work well. The policy preference for Q6b is "
         "statistically supported."),
        ("Single-item standards",
         "Nutritious Food (Q22) and Org Wellbeing (Q11a) are single items — reliability "
         "cannot be assessed. Consider whether they need companion items."),
    ]

    for title, text in issues:
        doc.add_heading(title, level=3)
        doc.add_paragraph(text)

    # ── Section 6: KMO ────────────────────────────────────────────
    if "KMO Sampling Adequacy" in results:
        doc.add_page_break()
        doc.add_heading('7. Sampling Adequacy (KMO)', level=1)
        doc.add_paragraph(
            "Kaiser-Meyer-Olkin (KMO) measures whether the data is suitable for factor analysis. "
            "Overall KMO should be > 0.60. Per-item KMO below 0.50 suggests that item doesn't "
            "share enough variance with others for factor analysis."
        )
        add_table_to_doc(doc, results["KMO Sampling Adequacy"])

    doc.save(path)
    print(f"  Saved: {path}")
    return path


def main():
    print("=" * 70)
    print("  GENERATING NOF ANALYSIS OUTPUTS (2024 data only)")
    print("=" * 70)

    print("\n  Extracting data from NSS24 Excel...")
    df = extract_all_nof_items()

    print("\n  Computing all results...")
    results = compute_all_results(df)

    print(f"\n  Generated {len(results)} result tables:")
    for name, tbl in results.items():
        print(f"    {name:40s}  {tbl.shape[0]} rows x {tbl.shape[1]} cols")

    print("\n  Writing outputs...")
    write_excel(results)
    write_csvs(results)
    write_word(results)

    print("\n  ALL DONE")
    print(f"  Outputs in: {os.path.abspath(OUTPUT_DIR)}/")


if __name__ == "__main__":
    main()
